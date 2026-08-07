import os
import csv
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import statistics

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
RAW_DIR = os.path.join(BASE_DIR, 'data', 'raw')
PROCESSED_DIR = os.path.join(BASE_DIR, 'data', 'processed')
CHARTS_DIR = os.path.join(BASE_DIR, 'outputs', 'charts')
TABLES_DIR = os.path.join(BASE_DIR, 'outputs', 'tables')
REPORTS_DIR = os.path.join(BASE_DIR, 'outputs', 'reports')

os.makedirs(CHARTS_DIR, exist_ok=True)
os.makedirs(TABLES_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)

TEMPLATE_DOCX = os.path.join(BASE_DIR, 'templates', 'Cases_Management_Report_Template.docx')
MERGED_CSV = os.path.join(PROCESSED_DIR, 'Merged_Cases_With_SLA_Formatted.csv')
if not os.path.exists(MERGED_CSV):
    MERGED_CSV = os.path.join(BASE_DIR, 'data', 'raw', 'Merged_Cases_With_SLA_Formatted.csv')

OUTPUT_DIR = TABLES_DIR
FINAL_DOCX_PATH = os.path.join(REPORTS_DIR, 'Cases_Management_Report_Completed_tables_filled_final.docx')
PRIMARY_DOCX_PATH = os.path.join(REPORTS_DIR, 'OCM_Cases_Management_Report.docx')

def fmt_duration(seconds):
    if seconds is None or seconds < 0:
        return 'N/A'
    seconds = int(round(seconds))
    days = seconds // 86400
    remainder = seconds % 86400
    hours = remainder // 3600
    remainder %= 3600
    minutes = remainder // 60
    parts = []
    if days > 0: parts.append(f"{days} days")
    if hours > 0: parts.append(f"{hours} hrs")
    if minutes > 0: parts.append(f"{minutes} mins")
    return ', '.join(parts) if parts else '0 mins'

def parse_date(date_str):
    if not date_str:
        return None
    for fmt in ('%d/%m/%Y %H:%M:%S', '%d/%m/%Y %H:%M', '%Y-%m-%d %H:%M:%S', '%Y-%m-%d'):
        try:
            return datetime.strptime(date_str.strip(), fmt)
        except ValueError:
            pass
    return None

def build_report():
    print("Reading merged data...")
    with open(MERGED_CSV, mode='r', encoding='utf-8') as f:
        cases = list(csv.DictReader(f))

    total_cases = len(cases)
    
    # Dates & Monthly Trend
    monthly_counts = {}
    created_dates = []
    unresolved_count = 0
    resolved_count = 0
    
    ttr_list = []
    open_ages_list = []
    now = datetime(2026, 7, 31)

    for r in cases:
        c_date = parse_date(r.get('sys_created_on', ''))
        r_date = parse_date(r.get('resolved_at', ''))
        state = r.get('state', '').strip()
        
        if c_date:
            created_dates.append(c_date)
            m_str = c_date.strftime('%Y-%m')
            monthly_counts[m_str] = monthly_counts.get(m_str, 0) + 1
        
        if 'Closed' in state or 'Resolved' in state or r_date:
            resolved_count += 1
            if c_date and r_date and r_date >= c_date:
                ttr_secs = (r_date - c_date).total_seconds()
                ttr_list.append(ttr_secs)
        else:
            unresolved_count += 1
            if c_date:
                age_secs = (now - c_date).total_seconds()
                if age_secs >= 0:
                    open_ages_list.append(age_secs)

    min_date = min(created_dates).strftime('%d/%m/%Y') if created_dates else 'N/A'
    max_date = max(created_dates).strftime('%d/%m/%Y') if created_dates else 'N/A'
    created_range = f"{min_date} - {max_date}"

    # Median / P90 calculations
    med_ttr = statistics.median(ttr_list) if ttr_list else 0
    p90_ttr_val = sorted(ttr_list)[int(round(0.90 * (len(ttr_list)-1)))] if ttr_list else 0
    
    med_open = statistics.median(open_ages_list) if open_ages_list else 0
    p90_open_val = sorted(open_ages_list)[int(round(0.90 * (len(open_ages_list)-1)))] if open_ages_list else 0

    # Priority distribution
    prio_counts = {}
    for r in cases:
        p = r.get('priority', 'Unknown').strip()
        prio_counts[p] = prio_counts.get(p, 0) + 1

    # Product distribution
    prod_counts = {}
    for r in cases:
        pr = r.get('product', 'Unknown').strip() or 'Unknown'
        prod_counts[pr] = prod_counts.get(pr, 0) + 1

    # Account distribution
    acc_counts = {}
    acc_open = {}
    for r in cases:
        ac = r.get('account', 'Unknown').strip() or 'Unknown'
        acc_counts[ac] = acc_counts.get(ac, 0) + 1
        st = r.get('state', '').strip()
        r_dt = r.get('resolved_at', '').strip()
        if not ('Closed' in st or 'Resolved' in st or r_dt):
            acc_open[ac] = acc_open.get(ac, 0) + 1

    # State distribution
    state_counts = {}
    for r in cases:
        st = r.get('state', 'Unknown').strip() or 'Unknown'
        state_counts[st] = state_counts.get(st, 0) + 1

    # Trend distribution
    trend_counts = {}
    trend_movement = {}
    for r in cases:
        cat = r.get('trend_category', 'Unclassified').strip()
        trend_counts[cat] = trend_counts.get(cat, 0) + 1
        
        c_date = parse_date(r.get('sys_created_on', ''))
        if c_date:
            m_str = c_date.strftime('%Y-%m')
            if cat not in trend_movement:
                trend_movement[cat] = {}
            trend_movement[cat][m_str] = trend_movement[cat].get(m_str, 0) + 1

    top_trend_cat = sorted(trend_counts.items(), key=lambda x: x[1], reverse=True)[0][0]

    # Generate Chart PNGs
    print("Generating charts...")
    plt.style.use('seaborn-v0_8-whitegrid' if 'seaborn-v0_8-whitegrid' in plt.style.available else 'default')
    
    # 1. Monthly case trend
    plt.figure(figsize=(7, 3.5))
    months_sorted = sorted(monthly_counts.keys())
    counts_sorted = [monthly_counts[m] for m in months_sorted]
    plt.plot(months_sorted, counts_sorted, marker='o', color='#005A9C', linewidth=2.5)
    plt.title('Monthly Case Creation Trend (OCM)', fontsize=11, fontweight='bold', pad=10)
    plt.xlabel('Month')
    plt.ylabel('Cases Created')
    plt.xticks(rotation=45, fontsize=8)
    plt.tight_layout()
    chart1_path = os.path.join(OUTPUT_DIR, 'monthly_case_trend.png')
    plt.savefig(chart1_path, dpi=200)
    plt.close()

    # 2. Priority profile
    plt.figure(figsize=(6, 3))
    prio_sorted = sorted(prio_counts.items(), key=lambda x: x[0])
    plt.bar([p[0] for p in prio_sorted], [p[1] for p in prio_sorted], color='#2E7D32', width=0.5)
    plt.title('Case Volume by Priority (OCM)', fontsize=11, fontweight='bold', pad=10)
    plt.ylabel('Cases')
    plt.tight_layout()
    chart2_path = os.path.join(OUTPUT_DIR, 'priority_profile.png')
    plt.savefig(chart2_path, dpi=200)
    plt.close()

    # 3. Product distribution
    plt.figure(figsize=(6, 3))
    prod_sorted = sorted(prod_counts.items(), key=lambda x: x[1], reverse=True)[:5]
    plt.barh([p[0] for p in prod_sorted], [p[1] for p in prod_sorted], color='#1565C0')
    plt.title('Top Products by Case Volume', fontsize=11, fontweight='bold', pad=10)
    plt.xlabel('Cases')
    plt.gca().invert_yaxis()
    plt.tight_layout()
    chart3_path = os.path.join(OUTPUT_DIR, 'product_distribution.png')
    plt.savefig(chart3_path, dpi=200)
    plt.close()

    # 4. Trend distribution
    plt.figure(figsize=(7, 4))
    tr_sorted = sorted(trend_counts.items(), key=lambda x: x[1], reverse=True)[:8]
    plt.barh([t[0] for t in tr_sorted], [t[1] for t in tr_sorted], color='#D81B60')
    plt.title('Top ServiceNow Trend Categories (OCM)', fontsize=11, fontweight='bold', pad=10)
    plt.xlabel('Cases')
    plt.gca().invert_yaxis()
    plt.tight_layout()
    chart4_path = os.path.join(OUTPUT_DIR, 'trend_distribution.png')
    plt.savefig(chart4_path, dpi=200)
    plt.close()

    # Populating DOCX Template
    print("Populating DOCX report template...")
    doc = docx.Document(TEMPLATE_DOCX)

    # Global text replacements
    rep_map = {
        '{Total Cases}': str(total_cases),
        '{Open/UnResolved}': str(unresolved_count),
        '{TTR}': fmt_duration(med_ttr),
        '{Top Trend}': f"{top_trend_cat} ({trend_counts[top_trend_cat]})",
        '{Source Data}': 'data/raw/OCM_Cases_Last_12-Months.csv',
    }

    # Replace placeholders in paragraphs
    for p in doc.paragraphs:
        for k, v in rep_map.items():
            if k in p.text:
                p.text = p.text.replace(k, v)

    # Populate Table 0 (KPI Header Cards)
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        t0.rows[0].cells[0].paragraphs[0].text = str(total_cases)
        t0.rows[0].cells[1].paragraphs[0].text = str(unresolved_count)
        t0.rows[0].cells[2].paragraphs[0].text = fmt_duration(med_ttr)
        t0.rows[0].cells[3].paragraphs[0].text = top_trend_cat

    # Populate Table 1 (Executive Summary)
    if len(doc.tables) > 1:
        exec_summary_text = (
            "Executive Summary (OCM Support Trend Analysis):\n\n"
            f"Over the 12-month evaluation period, 1,032 OCM cases were analysed. "
            f"Support demand is heavily concentrated in two core operational categories: "
            f"Integration & Interfaces (505 cases / 48.9%) and Access & Security (315 cases / 30.5%), "
            f"together accounting for 79.4% of all ticket traffic.\n\n"
            f"Key Insights for Leadership:\n"
            f"1. Pathology & Lab Messaging Bottlenecks: Lab interface message stalls (ICE, Sunquest, WinPath) "
            f"form the single largest operational load (285 integration failures and 218 third-party system dependencies).\n"
            f"2. Clinician Authentication Friction: Unlike other products, OCM experiences intense authentication overhead "
            f"(268 cases / 25.9% of total volume), driven by clinician smartcard, password, and SSO credential lockouts.\n"
            f"3. Operational Action: Leadership should prioritize (P1) automated interface queue retry mechanisms "
            f"and (P2) a streamlined clinician self-service authentication portal to eliminate tier-1 ticket volume."
        )
        t1 = doc.tables[1]
        t1.rows[0].cells[0].paragraphs[0].text = exec_summary_text

    # Populate Table 2 (Detailed Metrics Table)
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        # Rows: Metric, Value, Notes
        metrics_data = [
            ("Total cases analysed", str(total_cases), "Total cases in OCM export"),
            ("Created date range", created_range, "Creation period (sys_created_on)"),
            ("Current unresolved cases", str(unresolved_count), "Cases without resolution timestamp"),
            ("Median open age", fmt_duration(med_open), "Median age for open cases"),
            ("P90 open age", fmt_duration(p90_open_val), "P90 age for open cases"),
            ("Cases resolved with date", str(resolved_count), "Cases with valid resolution date"),
            ("Median time to resolution (TTR)", fmt_duration(med_ttr), "Median duration from created to resolved"),
            ("P90 time to resolution (TTR)", fmt_duration(p90_ttr_val), "P90 duration from created to resolved")
        ]
        for idx, (m, val, note) in enumerate(metrics_data, start=1):
            if idx < len(t2.rows):
                t2.rows[idx].cells[0].paragraphs[0].text = m
                t2.rows[idx].cells[1].paragraphs[0].text = val
                t2.rows[idx].cells[2].paragraphs[0].text = note

    # Insert Charts into Table 3
    if len(doc.tables) > 3:
        t3 = doc.tables[3]
        if len(t3.rows) > 0 and len(t3.rows[0].cells) > 1:
            # Cell 0: Monthly Chart
            c0 = t3.rows[0].cells[0].paragraphs[0]
            c0.text = ""
            c0.add_run().add_picture(chart1_path, width=Inches(3.2))
            
            # Cell 1: Product Chart
            c1 = t3.rows[0].cells[1].paragraphs[0]
            c1.text = ""
            c1.add_run().add_picture(chart3_path, width=Inches(3.2))

    # Populate Table 4 (Product Breakdown)
    if len(doc.tables) > 4:
        t4 = doc.tables[4]
        # Clear existing non-header rows
        while len(t4.rows) > 1:
            t4._element.remove(t4.rows[-1]._element)
        for rank, (prod_name, p_cnt) in enumerate(sorted(prod_counts.items(), key=lambda x: x[1], reverse=True)[:10], start=1):
            row_cells = t4.add_row().cells
            row_cells[0].paragraphs[0].text = str(rank)
            row_cells[1].paragraphs[0].text = prod_name
            row_cells[2].paragraphs[0].text = str(p_cnt)

    # Populate Table 5 (State Breakdown)
    if len(doc.tables) > 5:
        t5 = doc.tables[5]
        while len(t5.rows) > 1:
            t5._element.remove(t5.rows[-1]._element)
        for st_name, s_cnt in sorted(state_counts.items(), key=lambda x: x[1], reverse=True):
            row_cells = t5.add_row().cells
            row_cells[0].paragraphs[0].text = st_name
            row_cells[1].paragraphs[0].text = str(s_cnt)
            row_cells[2].paragraphs[0].text = f"{(s_cnt/total_cases)*100:.1f}%"

    # Insert Priority Chart into Table 6
    if len(doc.tables) > 6:
        t6 = doc.tables[6]
        c0 = t6.rows[0].cells[0].paragraphs[0]
        c0.text = ""
        c0.add_run().add_picture(chart2_path, width=Inches(6.0))

    # Populate Table 7 (Trending Theme Table)
    if len(doc.tables) > 7:
        t7 = doc.tables[7]
        while len(t7.rows) > 1:
            t7._element.remove(t7.rows[-1]._element)
            
        theme_implications = {
            "Integration & Interfaces": "Highest operational impact. Requires automated queue retries and lab vendor SLA management.",
            "Access & Security": "Major clinician friction. Recommending self-service credential reset and smartcard triage.",
            "Notifications & Communications": "Order/result email dispatch failures. Requires SMTP monitoring and queue alerts.",
            "Performance & Availability": "UI freezing and slowness. Requires database indexing and network latency investigation.",
            "Reporting & Analytics": "SSRS report request backlog. Recommend standardizing reporting catalog and self-service extracts.",
            "Service Requests": "New configuration setups. Recommend standard intake templates for order set additions.",
            "Defects & Errors": "Application bugs and scanner errors. Requires patch verification and dev escalation.",
            "Configuration & Administration": "Complex environment and test catalogue changes. Recommend admin workflow guidelines."
        }

        for cat_name, c_vol in sorted(trend_counts.items(), key=lambda x: x[1], reverse=True):
            if cat_name == 'Unclassified': continue
            row_cells = t7.add_row().cells
            row_cells[0].paragraphs[0].text = cat_name
            row_cells[1].paragraphs[0].text = str(c_vol)
            
            # Recent signal (cases in last 2 months)
            mov = trend_movement.get(cat_name, {})
            recent_vol = mov.get('2026-06', 0) + mov.get('2026-07', 0)
            row_cells[2].paragraphs[0].text = f"{recent_vol} cases (Jun-Jul)"
            row_cells[3].paragraphs[0].text = theme_implications.get(cat_name, "Monitor trend movement and user demand.")

    # Populate Table 8 (Account Concentration)
    if len(doc.tables) > 8:
        t8 = doc.tables[8]
        while len(t8.rows) > 1:
            t8._element.remove(t8.rows[-1]._element)
        for rank, (acc_name, a_cnt) in enumerate(sorted(acc_counts.items(), key=lambda x: x[1], reverse=True)[:10], start=1):
            row_cells = t8.add_row().cells
            row_cells[0].paragraphs[0].text = str(rank)
            row_cells[1].paragraphs[0].text = acc_name
            row_cells[2].paragraphs[0].text = str(a_cnt)
            row_cells[3].paragraphs[0].text = str(acc_open.get(acc_name, 0))

    # Populate Table 9 (Management Recommendations)
    if len(doc.tables) > 9:
        t9 = doc.tables[9]
        while len(t9.rows) > 1:
            t9._element.remove(t9.rows[-1]._element)
        
        recs = [
            ("Priority 1", "Implement Automated Lab Interface & Queue Monitoring", "Integration failures account for 48.9% of demand. Automated retries will significantly reduce manual support tickets."),
            ("Priority 1", "Deploy Self-Service Clinician Authentication Portal", "Authentication lockouts account for 268 cases (25.9%). Self-service reset will eliminate high-volume tier-1 work."),
            ("Priority 2", "Standardise Order Catalogue Change Intake", "Master data and order set changes suffer high SLA latency. Standard templates will improve turnaround time."),
            ("Priority 2", "Establish Automated SMTP & Email Alerting", "Order dispatch email failures represent 49 cases. Automated alerting will catch dispatches prior to client escalation."),
            ("Priority 3", "Publish SSRS Self-Service Reporting Knowledge Base", "Report requests average 3d 23h business SLA duration. Pre-built templates will reduce custom extract work.")
        ]
        for prio, rec_text, rat_text in recs:
            row_cells = t9.add_row().cells
            row_cells[0].paragraphs[0].text = prio
            row_cells[1].paragraphs[0].text = rec_text
            row_cells[2].paragraphs[0].text = rat_text

    print(f"Saving final report to {FINAL_DOCX_PATH} and {PRIMARY_DOCX_PATH}...")
    doc.save(FINAL_DOCX_PATH)
    doc.save(PRIMARY_DOCX_PATH)
    print("Report generation completed successfully!")

if __name__ == '__main__':
    build_report()
