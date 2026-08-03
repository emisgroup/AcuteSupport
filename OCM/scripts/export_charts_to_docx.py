from docx import Document
from docx.shared import Inches
import re
import os
import pandas as pd

# Base directories (dynamic)
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
RAW_DIR = os.path.join(BASE_DIR, 'data', 'raw')
charts_dir = os.path.join(BASE_DIR, 'outputs', 'charts')
tables_dir = os.path.join(BASE_DIR, 'outputs', 'tables')
reports_dir = os.path.join(BASE_DIR, 'outputs', 'reports')
os.makedirs(reports_dir, exist_ok=True)

output_docx = os.path.join(reports_dir, 'Cases_Management_Report_Completed.docx')
template = os.path.join(BASE_DIR, 'templates', 'Cases_Management_Report_Template.docx')

# Load KPI values
kpi_csv = os.path.join(tables_dir, 'kpi_overview.csv')
kpi = {}
if os.path.exists(kpi_csv):
    kdf = pd.read_csv(kpi_csv)
    if len(kdf)>0:
        kpi = kdf.iloc[0].to_dict()

# available charts and tables
charts = {
    'monthly_case_trend':'monthly_case_trend.png',
    'priority_profile':'priority_profile.png',
    'product_distribution':'product_distribution.png',
    'trend_distribution':'trend_distribution.png',
    'trend_movement':'trend_movement.png',
    'median_vs_p90':'median_vs_p90.png',
    'sla_performance':'sla_performance.png'
}

tables = {
    'monthly_case_trend':'monthly_case_trend.csv',
    'priority_profile':'priority_profile.csv',
    'product_distribution':'product_distribution.csv',
    'trend_distribution':'trend_distribution.csv',
    'trend_movement':'trend_movement.csv'
}

# helper: replace placeholder text in runs
placeholder_re = re.compile(r"\{([^}]+)\}")

# explicit mapping for custom placeholders in template
explicit_kpi_map = {

    'Total Cases': str(kpi.get('total_cases', '')),
    'Open/UnResolved': str(kpi.get('open_cases', '')),
    'TTR': str(kpi.get('median_ttr', '')),
    'Top Trend': 'Integration & Interfaces / Third-Party Systems'
}

explicit_chart_map = {
    'CASES CREATED BY MONTH CHART': 'monthly_case_trend.png',
    'Top Products by Case Volume Chart': 'product_distribution.png',
    'Priority Profile Chart': 'priority_profile.png'
}

def replace_placeholders_in_paragraph(paragraph):
    full_text = ''.join(run.text for run in paragraph.runs)
    matches = placeholder_re.findall(full_text)
    if not matches:
        return False
    new_text = full_text
    for key in matches:
        k = key.strip()
        if k in explicit_kpi_map:
            new_text = new_text.replace('{' + key + '}', explicit_kpi_map[k])
        elif k in kpi and kpi.get(k) not in [None, '']:
            new_text = new_text.replace('{' + key + '}', str(kpi.get(k)))
        elif k in kpi:
            new_text = new_text.replace('{' + key + '}', str(kpi.get(k)))
    for i in range(len(paragraph.runs)-1, -1, -1):
        paragraph.runs[i].clear()
    if new_text:
        paragraph.add_run(new_text)
    return True


def insert_image_after_paragraph(doc, paragraph, image_path, width_inches=6):
    p = paragraph._p
    new_p = doc.add_paragraph()
    p.getparent().insert(p.getparent().index(p)+1, new_p._p)
    run = new_p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return new_p


def insert_table_after_paragraph(doc, paragraph, csv_path, max_rows=10):
    if not os.path.exists(csv_path):
        return
    df = pd.read_csv(csv_path)
    rows = df.shape[0]
    cols = df.shape[1]
    new_p = doc.add_paragraph()
    paragraph._p.getparent().insert(paragraph._p.getparent().index(paragraph._p)+1, new_p._p)
    table = doc.add_table(rows=1, cols=cols)
    hdr_cells = table.rows[0].cells
    for ci, col in enumerate(df.columns):
        hdr_cells[ci].text = str(col)
    for r in range(min(max_rows, rows)):
        row_cells = table.add_row().cells
        for ci, col in enumerate(df.columns):
            val = df.iloc[r, ci]
            row_cells[ci].text = '' if pd.isna(val) else str(val)
    return table

# Load template
doc = Document(template)

# First pass: replace KPI placeholders in paragraphs
for para in doc.paragraphs:
    replace_placeholders_in_paragraph(para)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_placeholders_in_paragraph(para)

# Helper to replace explicit chart placeholders in a paragraph
def replace_explicit_charts(para):
    for ph, img_name in explicit_chart_map.items():
        placeholder_tag = '{' + ph + '}'
        if placeholder_tag in para.text:
            img_path = os.path.join(charts_dir, img_name)
            if os.path.exists(img_path):
                for run in para.runs:
                    run.text = ''
                run = para.add_run()
                run.add_picture(img_path, width=Inches(3.2 if 'Table' in str(type(para._p.getparent())) else 6))

for para in list(doc.paragraphs):
    replace_explicit_charts(para)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in list(cell.paragraphs):
                replace_explicit_charts(para)

chart_pattern = re.compile(r"\{\s*(?:Chart|Image)\s*:\s*([^}]+)\s*\}", re.IGNORECASE)


table_pattern = re.compile(r"\{\s*Table\s*:\s*([^}]+)\s*\}", re.IGNORECASE)

paras = list(doc.paragraphs)
for para in paras:
    text = para.text
    m = chart_pattern.search(text)
    if m:
        key = m.group(1).strip()
        img_file = charts.get(key, key + '.png')
        img_path = os.path.join(charts_dir, img_file)
        if os.path.exists(img_path):
            for run in para.runs:
                run.text = run.text.replace(m.group(0), '')
            insert_image_after_paragraph(doc, para, img_path)
    m2 = table_pattern.search(text)
    if m2:
        key = m2.group(1).strip()
        csv_file = tables.get(key, key + '.csv')
        csv_path = os.path.join(tables_dir, csv_file)
        if os.path.exists(csv_path):
            for run in para.runs:
                run.text = run.text.replace(m2.group(0), '')
            insert_table_after_paragraph(doc, para, csv_path)

simple_chart_pattern = re.compile(r"\{\s*([a-z0-9_]+)\s*\}", re.IGNORECASE)
for para in list(doc.paragraphs):
    text = para.text
    m = simple_chart_pattern.findall(text)
    if not m:
        continue
    for token in m:
        key = token.strip()
        if key.lower() in charts:
            img_path = os.path.join(charts_dir, charts[key.lower()])
            if os.path.exists(img_path):
                for run in para.runs:
                    run.text = run.text.replace('{' + token + '}', '')
                insert_image_after_paragraph(doc, para, img_path)


# Final: save completed doc
os.makedirs(reports_dir, exist_ok=True)
doc.save(output_docx)
print('Saved completed report to', output_docx)
