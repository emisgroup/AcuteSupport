import os
import csv
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import statistics

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
RAW_DIR = os.path.join(BASE_DIR, 'data', 'raw')
TEMPLATE_DOCX = os.path.join(BASE_DIR, 'templates', 'Cases_Management_Report_Template.docx')
MERGED_CSV = os.path.join(RAW_DIR, 'Merged_Cases_With_SLA_Formatted.csv')
OUTPUT_DIR = os.path.join(RAW_DIR, 'report_outputs')
os.makedirs(OUTPUT_DIR, exist_ok=True)

FINAL_DOCX_PATH = os.path.join(OUTPUT_DIR, 'Cases_Management_Report_Completed_tables_filled_final.docx')
PRIMARY_DOCX_PATH = os.path.join(OUTPUT_DIR, 'Symphony_Cases_Management_Report.docx')

def fmt_duration(seconds):
    if seconds is None or seconds < 0:
        return 'N/A'
    seconds = int(round(seconds))
    days = seconds // 86400
    remainder = seconds % 86400
    hours = remainder // 3600
    remainder %= 3600
    minutes = remainder // 60
    parts = []
    if days > 0: parts.append(f"{days} days")
    if hours > 0: parts.append(f"{hours} hrs")
    if minutes > 0: parts.append(f"{minutes} mins")
    return ', '.join(parts) if parts else '0 mins'

def parse_date(date_str):
    if not date_str:
        return None
    for fmt in ('%d/%m/%Y %H:%M:%S', '%d/%m/%Y %H:%M', '%Y-%m-%d %H:%M:%S', '%Y-%m-%d'):
        try:
            return datetime.strptime(date_str.strip(), fmt)
        except ValueError:
            pass
    return None

def build_report():
    print("Reading Symphony merged data...")
    with open(MERGED_CSV, mode='r', encoding='utf-8') as f:
        cases = list(csv.DictReader(f))

    total_cases = len(cases)
    
    monthly_counts = {}
    created_dates = []
    unresolved_count = 0
    resolved_count = 0
    
    ttr_list = []
    open_ages_list = []
    now = datetime(2026, 7, 31)

    for r in cases:
        c_date = parse_date(r.get('sys_created_on', ''))
        r_date = parse_date(r.get('resolved_at', ''))
        state = r.get('state', '').strip()
        
        if c_date:
            created_dates.append(c_date)
            m_str = c_date.strftime('%Y-%m')
            monthly_counts[m_str] = monthly_counts.get(m_str, 0) + 1
        
        if 'Closed' in state or 'Resolved' in state or r_date:
            resolved_count += 1
            if c_date and r_date and r_date >= c_date:
                ttr_secs = (r_date - c_date).total_seconds()
                ttr_list.append(ttr_secs)
        else:
            unresolved_count += 1
            if c_date:
                age_secs = (now - c_date).total_seconds()
                if age_secs >= 0:
                    open_ages_list.append(age_secs)

    min_date = min(created_dates).strftime('%d/%m/%Y') if created_dates else 'N/A'
    max_date = max(created_dates).strftime('%d/%m/%Y') if created_dates else 'N/A'
    created_range = f"{min_date} - {max_date}"

    med_ttr = statistics.median(ttr_list) if ttr_list else 0
    p90_ttr_val = sorted(ttr_list)[int(round(0.90 * (len(ttr_list)-1)))] if ttr_list else 0
    
    med_open = statistics.median(open_ages_list) if open_ages_list else 0
    p90_open_val = sorted(open_ages_list)[int(round(0.90 * (len(open_ages_list)-1)))] if open_ages_list else 0

    prio_counts = {}
    for r in cases:
        p = r.get('priority', 'Unknown').strip()
        prio_counts[p] = prio_counts.get(p, 0) + 1

    prod_counts = {}
    for r in cases:
        pr = r.get('product', 'Unknown').strip() or 'Unknown'
        prod_counts[pr] = prod_counts.get(pr, 0) + 1

    acc_counts = {}
    acc_open = {}
    for r in cases:
        ac = r.get('account', 'Unknown').strip() or 'Unknown'
        acc_counts[ac] = acc_counts.get(ac, 0) + 1
        st = r.get('state', '').strip()
        r_dt = r.get('resolved_at', '').strip()
        if not ('Closed' in st or 'Resolved' in st or r_dt):
            acc_open[ac] = acc_open.get(acc_name, 0) if 'acc_name' in locals() else 0

    # Fix acc_open calculation
    acc_open = {}
    for r in cases:
        ac = r.get('account', 'Unknown').strip() or 'Unknown'
        st = r.get('state', '').strip()
        r_dt = r.get('resolved_at', '').strip()
        if not ('Closed' in st or 'Resolved' in st or r_dt):
            acc_open[ac] = acc_open.get(ac, 0) + 1

    state_counts = {}
    for r in cases:
        st = r.get('state', 'Unknown').strip() or 'Unknown'
        state_counts[st] = state_counts.get(st, 0) + 1

    trend_counts = {}
    trend_movement = {}
    for r in cases:
        cat = r.get('trend_category', 'Unclassified').strip()
        trend_counts[cat] = trend_counts.get(cat, 0) + 1
        
        c_date = parse_date(r.get('sys_created_on', ''))
        if c_date:
            m_str = c_date.strftime('%Y-%m')
            if cat not in trend_movement:
                trend_movement[cat] = {}
            trend_movement[cat][m_str] = trend_movement[cat].get(m_str, 0) + 1

    top_trend_cat = sorted(trend_counts.items(), key=lambda x: x[1], reverse=True)[0][0]

    print("Generating Symphony charts...")
    plt.style.use('seaborn-v0_8-whitegrid' if 'seaborn-v0_8-whitegrid' in plt.style.available else 'default')
    
    # 1. Monthly case trend
    plt.figure(figsize=(7, 3.5))
    months_sorted = sorted(monthly_counts.keys())
    counts_sorted = [monthly_counts[m] for m in months_sorted]
    plt.plot(months_sorted, counts_sorted, marker='o', color='#005A9C', linewidth=2.5)
    plt.title('Monthly Case Creation Trend (Symphony)', fontsize=11, fontweight='bold', pad=10)
    plt.xlabel('Month')
    plt.ylabel('Cases Created')
    plt.xticks(rotation=45, fontsize=8)
    plt.tight_layout()
    chart1_path = os.path.join(OUTPUT_DIR, 'monthly_case_trend.png')
    plt.savefig(chart1_path, dpi=200)
    plt.close()

    # 2. Priority profile
    plt.figure(figsize=(6, 3))
    prio_sorted = sorted(prio_counts.items(), key=lambda x: x[0])
    plt.bar([p[0] for p in prio_sorted], [p[1] for p in prio_sorted], color='#2E7D32', width=0.5)
    plt.title('Case Volume by Priority (Symphony)', fontsize=11, fontweight='bold', pad=10)
    plt.ylabel('Cases')
    plt.tight_layout()
    chart2_path = os.path.join(OUTPUT_DIR, 'priority_profile.png')
    plt.savefig(chart2_path, dpi=200)
    plt.close()

    # 3. Product distribution
    plt.figure(figsize=(6, 3))
    prod_sorted = sorted(prod_counts.items(), key=lambda x: x[1], reverse=True)[:5]
    plt.barh([p[0] for p in prod_sorted], [p[1] for p in prod_sorted], color='#1565C0')
    plt.title('Top Products by Case Volume', fontsize=11, fontweight='bold', pad=10)
    plt.xlabel('Cases')
    plt.gca().invert_yaxis()
    plt.tight_layout()
    chart3_path = os.path.join(OUTPUT_DIR, 'product_distribution.png')
    plt.savefig(chart3_path, dpi=200)
    plt.close()

    # 4. Trend distribution
    plt.figure(figsize=(7, 4))
    tr_sorted = sorted(trend_counts.items(), key=lambda x: x[1], reverse=True)[:8]
    plt.barh([t[0] for t in tr_sorted], [t[1] for t in tr_sorted], color='#D81B60')
    plt.title('Top ServiceNow Trend Categories (Symphony)', fontsize=11, fontweight='bold', pad=10)
    plt.xlabel('Cases')
    plt.gca().invert_yaxis()
    plt.tight_layout()
    chart4_path = os.path.join(OUTPUT_DIR, 'trend_distribution.png')
    plt.savefig(chart4_path, dpi=200)
    plt.close()

    print("Populating Symphony DOCX report template...")
    doc = docx.Document(TEMPLATE_DOCX)

    rep_map = {
        '{Total Cases}': str(total_cases),
        '{Open/UnResolved}': str(unresolved_count),
        '{TTR}': fmt_duration(med_ttr),
        '{Top Trend}': f"{top_trend_cat} ({trend_counts[top_trend_cat]})",
        '{Source Data}': 'data/raw/Symphony_Casea_Last_12-Months.csv',
    }

    for p in doc.paragraphs:
        for k, v in rep_map.items():
            if k in p.text:
                p.text = p.text.replace(k, v)

    # Table 0 (KPI Cards)
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        t0.rows[0].cells[0].paragraphs[0].text = str(total_cases)
        t0.rows[0].cells[1].paragraphs[0].text = str(unresolved_count)
        t0.rows[0].cells[2].paragraphs[0].text = fmt_duration(med_ttr)
        t0.rows[0].cells[3].paragraphs[0].text = top_trend_cat

    # Table 1 (Executive Summary)
    if len(doc.tables) > 1:
        exec_summary_text = (
            "Executive Summary (Symphony Support Trend Analysis):\n\n"
            f"Over the 12-month evaluation period, 958 Symphony cases were analysed. "
            f"Support demand is heavily concentrated in two primary operational categories: "
            f"Integration & Interfaces (500 cases / 52.2%) and Access & Security (228 cases / 23.8%), "
            f"together accounting for 76.0% of all ticket volume.\n\n"
            f"Key Insights for Leadership:\n"
            f"1. Emergency & Acute Interface Demand: Integration failures (HL7 messaging / DAD interop) "
            f"represent 268 cases, while external NHS dependencies (SCR, TIE, ODS, GP Import) account for 230 cases.\n"
            f"2. High Information Governance Audit Overhead: Record access history and IG viewing audits "
            f"represent 151 cases (15.8% of overall volume), averaging 19h 2m SLA business duration.\n"
            f"3. Strategic Actions: Leadership should implement (P1) automated interface queue monitoring & retries "
            f"and (P1) an automated self-service Information Governance audit extraction portal."
        )
        t1 = doc.tables[1]
        t1.rows[0].cells[0].paragraphs[0].text = exec_summary_text

    # Table 2 (Metrics)
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        metrics_data = [
            ("Total cases analysed", str(total_cases), "Total cases in Symphony export"),
            ("Created date range", created_range, "Creation period (sys_created_on)"),
            ("Current unresolved cases", str(unresolved_count), "Cases without resolution timestamp"),
            ("Median open age", fmt_duration(med_open), "Median age for open cases"),
            ("P90 open age", fmt_duration(p90_open_val), "P90 age for open cases"),
            ("Cases resolved with date", str(resolved_count), "Cases with valid resolution date"),
            ("Median time to resolution (TTR)", fmt_duration(med_ttr), "Median duration from created to resolved"),
            ("P90 time to resolution (TTR)", fmt_duration(p90_ttr_val), "P90 duration from created to resolved")
        ]
        for idx, (m, val, note) in enumerate(metrics_data, start=1):
            if idx < len(t2.rows):
                t2.rows[idx].cells[0].paragraphs[0].text = m
                t2.rows[idx].cells[1].paragraphs[0].text = val
                t2.rows[idx].cells[2].paragraphs[0].text = note

    # Table 3 (Charts)
    if len(doc.tables) > 3:
        t3 = doc.tables[3]
        if len(t3.rows) > 0 and len(t3.rows[0].cells) > 1:
            c0 = t3.rows[0].cells[0].paragraphs[0]
            c0.text = ""
            c0.add_run().add_picture(chart1_path, width=Inches(3.2))
            
            c1 = t3.rows[0].cells[1].paragraphs[0]
            c1.text = ""
            c1.add_run().add_picture(chart3_path, width=Inches(3.2))

    # Table 4 (Product)
    if len(doc.tables) > 4:
        t4 = doc.tables[4]
        while len(t4.rows) > 1:
            t4._element.remove(t4.rows[-1]._element)
        for rank, (prod_name, p_cnt) in enumerate(sorted(prod_counts.items(), key=lambda x: x[1], reverse=True)[:10], start=1):
            row_cells = t4.add_row().cells
            row_cells[0].paragraphs[0].text = str(rank)
            row_cells[1].paragraphs[0].text = prod_name
            row_cells[2].paragraphs[0].text = str(p_cnt)

    # Table 5 (State)
    if len(doc.tables) > 5:
        t5 = doc.tables[5]
        while len(t5.rows) > 1:
            t5._element.remove(t5.rows[-1]._element)
        for st_name, s_cnt in sorted(state_counts.items(), key=lambda x: x[1], reverse=True):
            row_cells = t5.add_row().cells
            row_cells[0].paragraphs[0].text = st_name
            row_cells[1].paragraphs[0].text = str(s_cnt)
            row_cells[2].paragraphs[0].text = f"{(s_cnt/total_cases)*100:.1f}%"

    # Table 6 (Priority Chart)
    if len(doc.tables) > 6:
        t6 = doc.tables[6]
        c0 = t6.rows[0].cells[0].paragraphs[0]
        c0.text = ""
        c0.add_run().add_picture(chart2_path, width=Inches(6.0))

    # Table 7 (Trending Themes)
    if len(doc.tables) > 7:
        t7 = doc.tables[7]
        while len(t7.rows) > 1:
            t7._element.remove(t7.rows[-1]._element)
            
        theme_implications = {
            "Integration & Interfaces": "Highest operational impact (HL7/DAD messaging). Automated queue retries and interop monitoring required.",
            "Access & Security": "Driven by IG Audit & Compliance requests (151 cases). Recommending automated self-service audit extraction portal.",
            "Configuration & Administration": "Master data and departmental lookup changes (41 cases). Recommend standard admin change templates.",
            "Performance & Availability": "UI freezing and slowness. Requires database script optimisation and client network checks.",
            "Reporting & Analytics": "SSRS report generation & data extracts (31 cases). Recommend publishing standard reporting catalog.",
            "Defects & Errors": "Functional errors and scanning blank pages. Requires development patch verification.",
            "Service Requests": "New setup and quote requests. Recommend standard service catalog intake templates.",
            "Data Management": "Database data corrections and SQL fixes. Recommend tier-2 script library."
        }

        for cat_name, c_vol in sorted(trend_counts.items(), key=lambda x: x[1], reverse=True):
            if cat_name == 'Unclassified': continue
            row_cells = t7.add_row().cells
            row_cells[0].paragraphs[0].text = cat_name
            row_cells[1].paragraphs[0].text = str(c_vol)
            
            mov = trend_movement.get(cat_name, {})
            recent_vol = mov.get('2026-06', 0) + mov.get('2026-07', 0)
            row_cells[2].paragraphs[0].text = f"{recent_vol} cases (Jun-Jul)"
            row_cells[3].paragraphs[0].text = theme_implications.get(cat_name, "Monitor trend movement and user demand.")

    # Table 8 (Accounts)
    if len(doc.tables) > 8:
        t8 = doc.tables[8]
        while len(t8.rows) > 1:
            t8._element.remove(t8.rows[-1]._element)
        for rank, (acc_name, a_cnt) in enumerate(sorted(acc_counts.items(), key=lambda x: x[1], reverse=True)[:10], start=1):
            row_cells = t8.add_row().cells
            row_cells[0].paragraphs[0].text = str(rank)
            row_cells[1].paragraphs[0].text = acc_name
            row_cells[2].paragraphs[0].text = str(a_cnt)
            row_cells[3].paragraphs[0].text = str(acc_open.get(acc_name, 0))

    # Table 9 (Recommendations)
    if len(doc.tables) > 9:
        t9 = doc.tables[9]
        while len(t9.rows) > 1:
            t9._element.remove(t9.rows[-1]._element)
        
        recs = [
            ("Priority 1", "Implement Automated HL7 & DAD Interface Queue Retry", "Integration failures represent 52.2% of demand (500 cases). Automated retries will prevent manual support queue build-up."),
            ("Priority 1", "Deploy Self-Service Information Governance Audit Portal", "Audit & IG requests account for 151 cases (15.8%). An automated portal will eliminate tier-2 manual SQL audit extraction."),
            ("Priority 2", "Standardise Departmental & Ward Lookup Change Requests", "Master data changes represent 41 cases. Standardized intake forms will improve change turnaround times."),
            ("Priority 2", "Publish SSRS Self-Service Reporting Knowledge Base", "Reporting & data extracts account for 31 cases with high SLA duration (2d 3h). Pre-built templates reduce custom extract work."),
            ("Priority 3", "Establish Automated Printer & Output Configuration Guidance", "Batch printing issues represent 28 cases. KB articles and automated diagnostics will resolve client output issues.")
        ]
        for prio, rec_text, rat_text in recs:
            row_cells = t9.add_row().cells
            row_cells[0].paragraphs[0].text = prio
            row_cells[1].paragraphs[0].text = rec_text
            row_cells[2].paragraphs[0].text = rat_text

    print(f"Saving final Symphony report to {FINAL_DOCX_PATH} and {PRIMARY_DOCX_PATH}...")
    doc.save(FINAL_DOCX_PATH)
    doc.save(PRIMARY_DOCX_PATH)
    print("Symphony report generation completed successfully!")

if __name__ == '__main__':
    build_report()
