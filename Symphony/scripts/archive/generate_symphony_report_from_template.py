import re
from pathlib import Path
from collections import Counter
import pandas as pd
from docx import Document
from datetime import datetime

BASE = Path(r"C:\Users\lee.booth\Documents\02_ServiceNow\Management_Reports")
SYM_DIR = BASE / 'Symphony'
RAW = SYM_DIR / 'raw'
SCRIPTS = SYM_DIR / 'scripts'
REPORTS = SYM_DIR / 'reports'
TEMPLATES = SYM_DIR / 'templates'

CSV_FILE = RAW / 'Symphony_Cases_With_SLA_Last_12-Months.csv'
TEMPLATE = TEMPLATES / 'Symphony_Cases_Management_Report_Template.docx'

OUT_NAME = f"Symphony_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
OUT_PATH = REPORTS / OUT_NAME

# Helpers
re_day = re.compile(r"(\d+)\s*day", re.IGNORECASE)
re_hr = re.compile(r"(\d+)\s*hr", re.IGNORECASE)
re_min = re.compile(r"(\d+)\s*min", re.IGNORECASE)
re_sec = re.compile(r"(\d+)\s*sec", re.IGNORECASE)

STOPWORDS = set(["the","and","to","a","of","in","is","for","on","with","by","an","be","this","that","are","as","it","from","at","or","has","have"])


def parse_duration_to_seconds(s):
    if pd.isna(s) or s == "":
        return None
    try:
        if isinstance(s, (int,float)):
            return int(s)
        s_str = str(s)
        # if already digits
        if s_str.isdigit():
            return int(s_str)
    except Exception:
        pass
    days = hrs = mins = secs = 0
    m = re_day.search(s_str)
    if m: days = int(m.group(1))
    m = re_hr.search(s_str)
    if m: hrs = int(m.group(1))
    m = re_min.search(s_str)
    if m: mins = int(m.group(1))
    m = re_sec.search(s_str)
    if m: secs = int(m.group(1))
    total = days*86400 + hrs*3600 + mins*60 + secs
    return total if total>0 else 0


def seconds_to_human(seconds):
    if seconds is None or pd.isna(seconds):
        return ""
    seconds = int(seconds)
    parts = []
    days, rem = divmod(seconds, 86400)
    hrs, rem = divmod(rem, 3600)
    mins, secs = divmod(rem, 60)
    if days: parts.append(f"{days} day{'s' if days!=1 else ''}")
    if hrs: parts.append(f"{hrs} hr{'s' if hrs!=1 else ''}")
    if mins: parts.append(f"{mins} min{'s' if mins!=1 else ''}")
    if secs or not parts: parts.append(f"{secs} sec{'s' if secs!=1 else ''}")
    return ", ".join(parts)


# Read template headings
if not TEMPLATE.exists():
    raise SystemExit(f"Template not found: {TEMPLATE}")

doc_template = Document(TEMPLATE)
headings = [p.text.strip() for p in doc_template.paragraphs if p.style.name.lower().startswith('heading') and p.text.strip()]

# Read CSV
if not CSV_FILE.exists():
    raise SystemExit(f"CSV not found: {CSV_FILE}")

df = pd.read_csv(CSV_FILE, dtype=str)
# normalize
df.columns = [c.strip() for c in df.columns]
# keep Symphony only
if 'product' in df.columns:
    df = df[df['product'].fillna('').str.strip().str.lower()=='symphony']
# remove PRBs: where u_problem contains PRB or u_problem not empty
if 'u_problem' in df.columns:
    df = df[df['u_problem'].isna() | (df['u_problem'].str.strip()=='')]
# clean short desc
if 'short_description' in df.columns:
    df['short_description_clean'] = df['short_description'].fillna('').str.replace('\n',' ').str.replace('\r',' ').str.strip()
else:
    df['short_description_clean'] = ''
# parse dates
for col in ['sys_created_on','resolved_at']:
    if col in df.columns:
        df[col] = pd.to_datetime(df[col], dayfirst=True, errors='coerce')
# resolution
if 'sys_created_on' in df.columns and 'resolved_at' in df.columns:
    df['resolution_secs'] = (df['resolved_at'] - df['sys_created_on']).dt.total_seconds()
else:
    df['resolution_secs'] = None
# parse SLA durations
for col in ['SLA_business_duration','SLA_duration']:
    if col in df.columns:
        df[col + '_secs'] = df[col].apply(parse_duration_to_seconds)

# Metrics
total_cases = len(df)
p4_count = int(df['priority'].fillna('').str.startswith('4').sum()) if 'priority' in df.columns else 0
overall_avg_resolution = int(df['resolution_secs'].dropna().astype(float).mean()) if df['resolution_secs'].dropna().size>0 else None
overall_avg_sla_business = int(df['SLA_business_duration_secs'].dropna().astype(float).mean()) if 'SLA_business_duration_secs' in df.columns and df['SLA_business_duration_secs'].dropna().size>0 else None
overall_avg_sla_actual = int(df['SLA_duration_secs'].dropna().astype(float).mean()) if 'SLA_duration_secs' in df.columns and df['SLA_duration_secs'].dropna().size>0 else None

# Trends (top tokens)
tokens = []
for s in df['short_description_clean']:
    words = re.findall(r"\w+", str(s).lower())
    words = [w for w in words if w not in STOPWORDS and len(w)>2]
    tokens.extend(words)
common = Counter(tokens).most_common(6)

# Build report from template (append content matching headers)
doc = Document(TEMPLATE)

# Helper to append a heading and paragraphs
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading_and_paragraphs(doc, heading, paras):
    doc.add_heading(heading, level=2)
    for p in paras:
        doc.add_paragraph(p)

# Map of content by expected header names (case-insensitive)
content_map = {}
content_map['executive summary'] = [f"Total Symphony cases (last 12 months): {total_cases}", f"P4 case count: {p4_count}"]
if overall_avg_resolution:
    content_map['executive summary'].append(f"Average time to resolution: {seconds_to_human(overall_avg_resolution)}")
if overall_avg_sla_business:
    content_map['executive summary'].append(f"Average SLA business timing: {seconds_to_human(overall_avg_sla_business)}")
if overall_avg_sla_actual:
    content_map['executive summary'].append(f"Average SLA actual (SLA_Duration): {seconds_to_human(overall_avg_sla_actual)}")

# Trends
trend_paras = []
for tok, cnt in common:
    mask = df['short_description_clean'].str.lower().str.contains(tok)
    sub = df[mask]
    avg_res = int(sub['resolution_secs'].dropna().astype(float).mean()) if sub['resolution_secs'].dropna().size>0 else None
    avg_sla_bus = int(sub['SLA_business_duration_secs'].dropna().astype(float).mean()) if 'SLA_business_duration_secs' in sub.columns and sub['SLA_business_duration_secs'].dropna().size>0 else None
    avg_sla_act = int(sub['SLA_duration_secs'].dropna().astype(float).mean()) if 'SLA_duration_secs' in sub.columns and sub['SLA_duration_secs'].dropna().size>0 else None
    trend_paras.append(f"{tok.capitalize()} — {cnt} cases; Avg resolution: {seconds_to_human(avg_res)}; Avg SLA business: {seconds_to_human(avg_sla_bus)}; Avg SLA actual: {seconds_to_human(avg_sla_act)}")
content_map['trends'] = trend_paras

# SLA timings section
sla_paras = [f"Average SLA business timing: {seconds_to_human(overall_avg_sla_business)}", f"Average SLA actual (SLA_Duration): {seconds_to_human(overall_avg_sla_actual)}"]
content_map['sla timings'] = sla_paras

# Management recommendations (basic)
rec_paras = [
    "Create knowledgebase articles for top trends to reduce repeat cases.",
    "Introduce quick-reference guides and targeted training for common issues.",
    "Automate routine fixes and add monitoring for early detection.",
]
content_map['management recommendations'] = rec_paras

# P4 quick wins
p4 = df[df['priority'].fillna('').str.startswith('4')]
p4_paras = [f"P4 count: {len(p4)}"]
if len(p4)>0:
    top_p4 = Counter(p4['short_description_clean'].fillna('')).most_common(5)
    p4_paras.append('Top P4 summaries:')
    for k,cnt in top_p4:
        p4_paras.append(f"- {k} ({cnt})")
    p4_paras.append('Quick wins: Triage into fixes, document remediation, add alerts for recurrence.')
content_map['p4 quick wins'] = p4_paras

# For each heading found in template, append corresponding content (matching by keywords)
lower_headings = [h.lower() for h in headings]
# Preferred header order to append if present
preferred = ['executive summary','trends','sla timings','management recommendations','p4 quick wins']
for ph in preferred:
    # find if any template heading contains this key
    matches = [h for h in lower_headings if ph in h]
    if matches:
        # append using the original heading text from template for consistency
        orig = next(h for h in headings if ph in h.lower())
        add_heading_and_paragraphs(doc, orig, content_map.get(ph, ['No data']))

# If template did not contain some expected headers, append them at end
for ph in preferred:
    if not any(ph in h for h in lower_headings):
        add_heading_and_paragraphs(doc, ph.title(), content_map.get(ph, ['No data']))

# Save
REPORTS.mkdir(parents=True, exist_ok=True)
doc.save(OUT_PATH)
print(f"Report written to: {OUT_PATH}")
