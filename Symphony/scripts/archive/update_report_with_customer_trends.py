from pathlib import Path
from collections import Counter
import pandas as pd
from docx import Document
from datetime import datetime
import re

BASE = Path(r"C:\Users\lee.booth\Documents\02_ServiceNow\Management_Reports")
SYM = BASE / 'Symphony'
RAW = SYM / 'raw'
REPORTS = SYM / 'reports'

# Find most recent filled report (Symphony_Report_Filled_*)
reports = sorted(REPORTS.glob('Symphony_Report_Filled_*.docx'), key=lambda p: p.stat().st_mtime, reverse=True)
if not reports:
    raise SystemExit('No filled report found in reports/ matching Symphony_Report_Filled_*.docx')
IN_REPORT = reports[0]
OUT_REPORT = REPORTS / f"{IN_REPORT.stem}_customers_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

CSV = RAW / 'Symphony_Cases_With_SLA_Last_12-Months.csv'
if not CSV.exists():
    raise SystemExit('CSV not found: ' + str(CSV))

# Helpers
STOPWORDS = set(["the","and","to","a","of","in","is","for","on","with","by","an","be","this","that","are","as","it","from","at","or","has","have"]) 

def extract_tokens(text):
    words = re.findall(r"\w+", str(text).lower())
    words = [w for w in words if w not in STOPWORDS and len(w)>2]
    return words

# Read data
df = pd.read_csv(CSV, dtype=str)
df.columns = [c.strip() for c in df.columns]
# filter Symphony
if 'product' in df.columns:
    df = df[df['product'].fillna('').str.strip().str.lower()=='symphony']
# remove PRBs
if 'u_problem' in df.columns:
    df = df[df['u_problem'].isna() | (df['u_problem'].str.strip()=='')]
# clean short description
if 'short_description' in df.columns:
    df['short_description_clean'] = df['short_description'].fillna('').str.replace('\n',' ').str.replace('\r',' ').str.strip()
else:
    df['short_description_clean'] = ''

# compute basics
total_cases = len(df)
open_count = int(df[~df['state'].fillna('').str.lower().isin(['closed','resolved'])].shape[0]) if 'state' in df.columns else 0
p4_count = int(df['priority'].fillna('').str.startswith('4').sum()) if 'priority' in df.columns else 0

# top customers
top_customers = [c for c,_ in Counter(df['account'].fillna('Unknown')).most_common(10)]

# For each customer, find top 3 themes (tokens)
customer_themes = {}
for cust in top_customers:
    sub = df[df['account'].fillna('Unknown')==cust]
    tokens = []
    for s in sub['short_description_clean']:
        tokens.extend(extract_tokens(s))
    top = Counter(tokens).most_common(3)
    customer_themes[cust] = top

# Build executive summary text
summary_lines = []
summary_lines.append(f"Total Symphony cases (last 12 months): {total_cases}")
summary_lines.append(f"Open / unresolved: {open_count}")
summary_lines.append(f"P4 cases: {p4_count}")
# top 3 customers
top3 = Counter(df['account'].fillna('Unknown')).most_common(3)
if top3:
    summary_lines.append('Top customers by case volume: ' + ', '.join([f"{t[0]} ({t[1]})" for t in top3]))

# Load report
doc = Document(IN_REPORT)

# Fill executive summary table (table index 1 cell 0,0)
try:
    t1 = doc.tables[1]
    # Replace text with paragraph lines
    t1.rows[0].cells[0].text = '\n'.join(summary_lines)
except IndexError:
    # Append a heading and paragraph if table not present
    doc.add_heading('Executive summary', level=1)
    for line in summary_lines:
        doc.add_paragraph(line)

# Fill account concentration table (table index 9): replace 'Open' column with top themes
try:
    t9 = doc.tables[9]
    # Header assumed: Rank | Account | Cases | Open
    maxr = len(t9.rows)
    for i, cust in enumerate(top_customers, start=1):
        if i >= maxr: break
        row = t9.rows[i]
        cnt = int(df['account'].fillna('Unknown').value_counts().get(cust, 0))
        row.cells[0].text = str(i)
        row.cells[1].text = cust
        row.cells[2].text = str(cnt)
        themes = customer_themes.get(cust, [])
        theme_str = ', '.join([f"{t[0]} ({t[1]})" for t in themes]) if themes else ''
        # Put themes into 4th cell
        row.cells[3].text = theme_str
except IndexError:
    # append new table after last element
    doc.add_heading('Customer trends (top customers)', level=2)
    tbl = doc.add_table(rows=1, cols=4)
    hdr = tbl.rows[0].cells
    hdr[0].text = 'Rank'
    hdr[1].text = 'Account'
    hdr[2].text = 'Cases'
    hdr[3].text = 'Top themes (count)'
    for i, cust in enumerate(top_customers, start=1):
        row = tbl.add_row().cells
        row[0].text = str(i)
        row[1].text = cust
        cnt = int(df['account'].fillna('Unknown').value_counts().get(cust, 0))
        row[2].text = str(cnt)
        themes = customer_themes.get(cust, [])
        theme_str = ', '.join([f"{t[0]} ({t[1]})" for t in themes]) if themes else ''
        row[3].text = theme_str

# Save new report
doc.save(OUT_REPORT)
print('Updated report saved to:', OUT_REPORT)
