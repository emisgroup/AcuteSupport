from docx import Document
from docx.shared import Inches
import re
import os
import pandas as pd

# Base directories (dynamic)
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
RAW_DIR = os.path.join(BASE_DIR, 'data', 'raw')
charts_dir = os.path.join(BASE_DIR, 'outputs', 'charts')
tables_dir = os.path.join(BASE_DIR, 'outputs', 'tables')
reports_dir = os.path.join(BASE_DIR, 'outputs', 'reports')
os.makedirs(reports_dir, exist_ok=True)

output_docx = os.path.join(reports_dir, 'Cases_Management_Report_Completed.docx')
template = os.path.join(BASE_DIR, 'templates', 'Cases_Management_Report_Template.docx')

# Load KPI values
kpi_csv = os.path.join(tables_dir, 'kpi_overview.csv')
kpi = {}
if os.path.exists(kpi_csv):
    kdf = pd.read_csv(kpi_csv)
    if len(kdf)>0:
        kpi = kdf.iloc[0].to_dict()

# available charts and tables
charts = {
    'monthly_case_trend':'monthly_case_trend.png',
    'priority_profile':'priority_profile.png',
    'product_distribution':'product_distribution.png',
    'trend_distribution':'trend_distribution.png',
    'trend_movement':'trend_movement.png',
    'median_vs_p90':'median_vs_p90.png',
    'sla_performance':'sla_performance.png'
}

tables = {
    'monthly_case_trend':'monthly_case_trend.csv',
    'priority_profile':'priority_profile.csv',
    'product_distribution':'product_distribution.csv',
    'trend_distribution':'trend_distribution.csv',
    'trend_movement':'trend_movement.csv'
}

# helper: replace placeholder text in runs
placeholder_re = re.compile(r"\{([^}]+)\}")

def replace_placeholders_in_paragraph(paragraph):
    full_text = ''.join(run.text for run in paragraph.runs)
    matches = placeholder_re.findall(full_text)
    if not matches:
        return False
    new_text = full_text
    for key in matches:
        k = key.strip()
        # KPI replacement
        if k in kpi and kpi.get(k) not in [None, '']:
            new_text = new_text.replace('{' + key + '}', str(kpi.get(k)))
        # duration formatting fields may exist as *_formatted in kpi
        elif k in kpi:
            new_text = new_text.replace('{' + key + '}', str(kpi.get(k)))
        else:
            # leave as-is for chart/table placeholders; will be handled separately
            pass
    # set paragraph text to new_text (clear runs)
    for i in range(len(paragraph.runs)-1, -1, -1):
        paragraph.runs[i].clear()
    if new_text:
        paragraph.add_run(new_text)
    return True


def insert_image_after_paragraph(doc, paragraph, image_path, width_inches=6):
    p = paragraph._p
    idx = doc.paragraphs.index(paragraph)
    # Insert a new paragraph after this one
    new_p = doc.add_paragraph()
    # Move the new paragraph to after the target by reordering xml
    p.getparent().insert(p.getparent().index(p)+1, new_p._p)
    run = new_p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return new_p


def insert_table_after_paragraph(doc, paragraph, csv_path, max_rows=10):
    if not os.path.exists(csv_path):
        return
    df = pd.read_csv(csv_path)
    rows = df.shape[0]
    cols = df.shape[1]
    # build table with header + up to max_rows
    p = paragraph._p
    new_p = doc.add_paragraph()
    p.getparent().insert(p.getparent().index(p)+1, new_p._p)
    table = doc.add_table(rows=1, cols=cols)
    hdr_cells = table.rows[0].cells
    for ci, col in enumerate(df.columns):
        hdr_cells[ci].text = str(col)
    for r in range(min(max_rows, rows)):
        row_cells = table.add_row().cells
        for ci, col in enumerate(df.columns):
            val = df.iloc[r, ci]
            row_cells[ci].text = '' if pd.isna(val) else str(val)
    return table

# Load template
doc = Document(template)

# First pass: replace KPI placeholders in paragraphs
for para in doc.paragraphs:
    replace_placeholders_in_paragraph(para)

# Also check tables for placeholders in cell text
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_placeholders_in_paragraph(para)

# Second pass: handle chart and table placeholders specially
# Look for placeholders like {Chart:monthly_case_trend} or {Table:monthly_case_trend}
chart_pattern = re.compile(r"\{\s*(?:Chart|Image)\s*:\s*([^}]+)\s*\}", re.IGNORECASE)
table_pattern = re.compile(r"\{\s*Table\s*:\s*([^}]+)\s*\}", re.IGNORECASE)

# Walk paragraphs to insert images/tables
paras = list(doc.paragraphs)
for para in paras:
    text = para.text
    m = chart_pattern.search(text)
    if m:
        key = m.group(1).strip()
        img_file = charts.get(key, key + '.png')
        img_path = os.path.join(charts_dir, img_file)
        if os.path.exists(img_path):
            # remove placeholder text
            for run in para.runs:
                run.text = run.text.replace(m.group(0), '')
            insert_image_after_paragraph(doc, para, img_path)
    m2 = table_pattern.search(text)
    if m2:
        key = m2.group(1).strip()
        csv_file = tables.get(key, key + '.csv')
        csv_path = os.path.join(tables_dir, csv_file)
        if os.path.exists(csv_path):
            for run in para.runs:
                run.text = run.text.replace(m2.group(0), '')
            insert_table_after_paragraph(doc, para, csv_path)

# Additionally, simple placeholders matching chart names in braces will be replaced by images
simple_chart_pattern = re.compile(r"\{\s*([a-z0-9_]+)\s*\}", re.IGNORECASE)
for para in list(doc.paragraphs):
    text = para.text
    m = simple_chart_pattern.findall(text)
    if not m:
        continue
    for token in m:
        key = token.strip()
        if key.lower() in charts:
            img_path = os.path.join(charts_dir, charts[key.lower()])
            if os.path.exists(img_path):
                # replace token text and insert image
                for run in para.runs:
                    run.text = run.text.replace('{' + token + '}', '')
                insert_image_after_paragraph(doc, para, img_path)

# Final: save completed doc
os.makedirs(reports_dir, exist_ok=True)
doc.save(output_docx)
print('Saved completed report to', output_docx)
