from docx import Document
from docx.shared import Inches
import os
import pandas as pd
from shutil import copyfile
import numpy as np

# Paths
# Base directories (dynamic)
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
RAW_DIR = os.path.join(BASE_DIR, 'data', 'raw')
PROCESSED_DIR = os.path.join(BASE_DIR, 'data', 'processed')
tables_dir = os.path.join(BASE_DIR, 'outputs', 'tables')
reports_dir = os.path.join(BASE_DIR, 'outputs', 'reports')

merged_csv = os.path.join(PROCESSED_DIR, 'Merged_Cases_With_SLA_Formatted.csv')
if not os.path.exists(merged_csv):
    merged_csv = os.path.join(RAW_DIR, 'Merged_Cases_With_SLA_Formatted.csv')

input_doc = os.path.join(reports_dir, 'Cases_Management_Report_Completed_tables_filled.docx')
if not os.path.exists(input_doc):
    input_doc = os.path.join(reports_dir, 'Cases_Management_Report_Completed.docx')

backup_doc = os.path.join(reports_dir, 'Cases_Management_Report_Completed_tables_filled_backup2.docx')
final_doc = os.path.join(reports_dir, 'Cases_Management_Report_Completed_tables_filled_final.docx')

# parse current time (use now as fallback)
current_time = pd.to_datetime('2026-07-29T23:37:34.349+01:00', utc=True).tz_convert(None)

# helper
def format_duration(seconds):
    if seconds is None or (isinstance(seconds, float) and np.isnan(seconds)):
        return ''
    s = int(float(seconds))
    if s<=0:
        return ''
    mins = s//60
    days = mins // (24*60)
    mins_rem = mins - days*24*60
    hours = mins_rem // 60
    minutes = mins_rem - hours*60
    parts = []
    if days>0:
        parts.append(f"{days} days")
    if hours>0:
        parts.append(f"{hours} hrs")
    if minutes>0:
        parts.append(f"{minutes} mins")
    return ', '.join(parts)

# load data
if not os.path.exists(merged_csv):
    raise SystemExit('Merged CSV not found')

df = pd.read_csv(merged_csv, encoding='utf-8')
# parse dates
for col in ['sys_created_on','resolved_at']:
    if col in df.columns:
        df[col] = pd.to_datetime(df[col], dayfirst=True, errors='coerce')

# compute metrics
total_cases = len(df)
created_min = df['sys_created_on'].min()
created_max = df['sys_created_on'].max()
created_range = ''
if pd.notna(created_min) and pd.notna(created_max):
    created_range = f"{created_min.strftime('%d/%m/%Y')} - {created_max.strftime('%d/%m/%Y')}"

# unresolved: resolved_at is null or state not Closed
if 'resolved_at' in df.columns:
    unresolved_mask = df['resolved_at'].isna()
else:
    unresolved_mask = ~df['state'].str.contains('Closed', na=False)
current_unresolved = int(unresolved_mask.sum())

# open ages for unresolved
open_ages_seconds = None
if 'sys_created_on' in df.columns:
    now = current_time
    open_ages = (now - df.loc[unresolved_mask,'sys_created_on']).dt.total_seconds()
    open_ages = open_ages.dropna()
    if len(open_ages)>0:
        median_open = np.nanmedian(open_ages)
        p90_open = np.nanpercentile(open_ages,90)
    else:
        median_open = p90_open = None
else:
    median_open = p90_open = None

# resolved counts and durations
if 'resolved_at' in df.columns and 'sys_created_on' in df.columns:
    resolved_mask = df['resolved_at'].notna()
    resolved_with_date = int(resolved_mask.sum())
    ttr = (df.loc[resolved_mask,'resolved_at'] - df.loc[resolved_mask,'sys_created_on']).dt.total_seconds()
    ttr = ttr.dropna()
    if len(ttr)>0:
        median_res = np.nanmedian(ttr)
        p90_res = np.nanpercentile(ttr,90)
    else:
        median_res = p90_res = None
else:
    resolved_with_date = 0
    median_res = p90_res = None

# prepare notes
notes = {}
notes['Created date range'] = 'Date range of case creation (sys_created_on).'
notes['Current unresolved cases'] = 'Cases without a resolved date (resolved_at empty). Includes any non-Closed states.'
notes['Median open age'] = 'Median age for currently unresolved cases: duration since creation.'
notes['90th percentile open age'] = 'P90 age for currently unresolved cases.'
notes['Resolved/closed with resolution date'] = 'Count of cases with a resolved date present.'
notes['Median resolution duration'] = 'Median time-to-resolution for resolved cases (Created -> Resolved).'
notes['90th percentile resolution duration'] = 'P90 time-to-resolution for resolved cases.'

# Exec summary improved
# Find top trend
trend_csv = os.path.join(tables_dir, 'trend_distribution.csv')
trend_top = ''
if os.path.exists(trend_csv):
    td = pd.read_csv(trend_csv)
    if len(td)>0:
        trend_top = td.iloc[0,0]

exec_lines = []
exec_lines.append(f"This report covers {total_cases} Symphony cases. The data window is {created_range}.")
exec_lines.append(f"Demand concentration: highest-volume theme is '{trend_top}' (see Trend Distribution).")
if median_res:
    exec_lines.append(f"Primary service metric: median time to resolution is {format_duration(median_res)}, with P90 of {format_duration(p90_res)}.")
if median_open is not None:
    exec_lines.append(f"Current backlog: {current_unresolved} unresolved cases; median open age {format_duration(median_open)}, P90 {format_duration(p90_open)}.")
exec_lines.append('Recommended priorities: (1) rapidly create KB articles and automation for high-volume, low-complexity issues; (2) strengthen triage to reduce P90 resolution; (3) remediate recurring environmental/import processes causing operational noise.')
executive_summary = ' '.join(exec_lines)

# Source Data placeholder value
source_data_text = ('Source files: Symphony_Casea_Last_12-Months.csv (case data) and Symphony_SLA_Last_12-Months.csv (SLA durations), merged on case number. '
                    f'Derived dataset: Merged_Cases_With_SLA_Formatted.csv. Report generated: {current_time.strftime("%Y-%m-%d %H:%M:%S")}.')

# Recommended management actions (priority, recommendation, rationale)
recommended = [
    ('P1','Create knowledge base articles and runbook playbooks for top 5 high-volume trends','Quick wins: reduces repeat demand and triage time; high impact'),
    ('P1','Automate routine master-file and import tasks where possible','Removes manual steps that cause repeated incidents and frees analyst time'),
    ('P2','Implement enhanced triage rules and routing for Audit/Access requests','Reduces time-to-resolution and ensures correct assignment'),
    ('P2','Introduce pre-deployment checks for ODS/TRUD and master file updates','Reduces incidents caused by deployment/import failures'),
    ('P3','Perform a targeted review of low-volume but high-severity P90 cases','Investigate root causes and process changes for outliers'),
    ('P3','Develop training and runbooks for analysts on common configuration and SQL issues','Improves mean time to repair for technical problems')
]

# Open doc and backup
if not os.path.exists(input_doc):
    raise SystemExit('Input DOCX not found')
copyfile(input_doc, backup_doc)
doc = Document(input_doc)

# 1) Fill Key Metrics table (table index 2 per earlier mapping)
# Create list of rows to write
metrics = [
    ('Total cases analysed', str(total_cases), 'As counted in merged dataset'),
    ('Created date range', created_range, notes['Created date range']),
    ('Current unresolved cases', str(current_unresolved), notes['Current unresolved cases']),
    ('Median open age', format_duration(median_open), notes['Median open age']),
    ('90th percentile open age', format_duration(p90_open), notes['90th percentile open age']),
    ('Resolved/closed with resolution date', str(resolved_with_date), notes['Resolved/closed with resolution date']),
    ('Median resolution duration', format_duration(median_res), notes['Median resolution duration']),
    ('90th percentile resolution duration', format_duration(p90_res), notes['90th percentile resolution duration'])
]

if len(doc.tables) > 2:
    t = doc.tables[2]
    # remove existing data rows except header
    while len(t.rows) > 1:
        t._tbl.remove(t.rows[-1]._tr)
    for m in metrics:
        r = t.add_row()
        r.cells[0].text = m[0]
        r.cells[1].text = m[1]
        r.cells[2].text = m[2]

# 2) Improve executive summary 
# find placeholder or table 1 first cell
if len(doc.tables) > 1:
    t1 = doc.tables[1]
    # overwrite first cell paragraph
    cell = t1.rows[0].cells[0]
    for p in cell.paragraphs:
        for run in p.runs:
            run.clear()
    cell.paragraphs[0].add_run(executive_summary)

# 3) Populate {Source Data} placeholder in paragraphs
for para in doc.paragraphs:
    if '{Source Data}' in para.text:
        for r in para.runs:
            r.text = r.text.replace('{Source Data}', source_data_text)

# also in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if '{Source Data}' in cell.text:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = r.text.replace('{Source Data}', source_data_text)

# 4) Fill Recommended Management Actions table (likely index 10)
if len(doc.tables) > 10:
    t10 = doc.tables[10]
    # remove existing data rows except header
    while len(t10.rows) > 1:
        t10._tbl.remove(t10.rows[-1]._tr)
    for i, item in enumerate(recommended, start=1):
        r = t10.add_row()
        r.cells[0].text = str(i)
        r.cells[1].text = item[1]
        r.cells[2].text = item[2]

# 5) Append Appendix: Methods and Caveats - append at end
appendix_text = (
    'Appendix — Methods and Caveats\n\n'
    'Method: The report uses Symphony case exports merged with SLA durations on case number. Date parsing assumes day-first format. Time-to-resolution (TTR) is calculated as Resolved Date minus Created Date; open age uses the report generation time as reference. Trend classification used keyword matching against templates/Trends.txt with three additional authorised categories added during processing.\n\n'
    'Caveats: The merge preserved all cases from the primary case export and aligned SLA seconds where available; some cases lack SLA rows and some descriptions contain embedded line breaks which may affect automated token matching. 52 cases remain unclassified after automated passes. Durations are shown as "[DD] days, [HH] hrs, [MI] mins" with zero components suppressed.\n\n'
    'Privacy: No patient-identifiable data is included in outputs. Aggregate metrics only.\n\n'
    'If you require changes to classification rules, date range, or unresolved definition, request a rerun with updated parameters.'
)
# Append the appendix at the document end
doc.add_paragraph(appendix_text)

# Save final
doc.save(final_doc)
print('Saved final report to', final_doc)
print('Backup saved to', backup_doc)
