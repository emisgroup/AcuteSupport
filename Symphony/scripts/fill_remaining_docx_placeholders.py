from docx import Document
from docx.shared import Inches
import os
import pandas as pd
import math

# Base directories (dynamic)
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
RAW_DIR = os.path.join(BASE_DIR, 'data', 'raw')
out_dir = os.path.join(RAW_DIR, 'report_outputs')
docx_path = os.path.join(out_dir, 'Cases_Management_Report_Completed.docx')
backup_path = os.path.join(out_dir, 'Cases_Management_Report_Completed_before_fill.docx')


# load KPI
kpi_csv = os.path.join(out_dir,'kpi_overview.csv')
kpi = {}
if os.path.exists(kpi_csv):
    kdf = pd.read_csv(kpi_csv)
    if len(kdf)>0:
        kpi = kdf.iloc[0].to_dict()

# helper to format seconds to required string
def format_duration(seconds):
    try:
        if seconds=='' or pd.isna(seconds):
            return ''
        s = int(float(seconds))
        if s<=0:
            return ''
        minutes = s//60
        days = minutes // (24*60)
        minutes_rem = minutes - days*24*60
        hours = minutes_rem // 60
        mins = minutes_rem - hours*60
        parts = []
        if days>0:
            parts.append(f"{days} days")
        if hours>0:
            parts.append(f"{hours} hrs")
        if mins>0:
            parts.append(f"{mins} mins")
        return ', '.join(parts)
    except Exception:
        return ''

# read supporting tables
trend_dist_path = os.path.join(out_dir,'trend_distribution.csv')
monthly_chart = os.path.join(out_dir,'monthly_case_trend.png')
product_chart = os.path.join(out_dir,'product_distribution.png')
priority_chart = os.path.join(out_dir,'priority_profile.png')
trend_chart = os.path.join(out_dir,'trend_distribution.png')
trend_movement_chart = os.path.join(out_dir,'trend_movement.png')
median_vs_p90_chart = os.path.join(out_dir,'median_vs_p90.png')
sla_chart = os.path.join(out_dir,'sla_performance.png')

trend_top = ''
if os.path.exists(trend_dist_path):
    td = pd.read_csv(trend_dist_path)
    if len(td)>0:
        trend_top = td.iloc[0,0]

# Create executive summary (short, management-readable)
# Use kpi values if available
total_cases = int(kpi.get('total_cases', 0) if kpi.get('total_cases')!='' else 0)
closed_cases = int(kpi.get('closed_cases', 0) if kpi.get('closed_cases')!='' else 0)
open_cases = int(kpi.get('open_cases', 0) if kpi.get('open_cases')!='' else 0)
median_ttr_s = kpi.get('median_ttr_seconds', '')
p90_ttr_s = kpi.get('p90_ttr_seconds', '')
median_ttr = format_duration(median_ttr_s)
p90_ttr = format_duration(p90_ttr_s)
avg_sla = kpi.get('avg_sla_business', '')

exec_summary = []
exec_summary.append(f"Total cases in period: {total_cases}.")
if open_cases:
    exec_summary.append(f"Open cases: {open_cases}; Closed: {closed_cases}.")
if median_ttr:
    exec_summary.append(f"Median time to resolution: {median_ttr}.")
if p90_ttr:
    exec_summary.append(f"P90 time to resolution: {p90_ttr}.")
if trend_top:
    exec_summary.append(f"Top trend by volume: {trend_top}.")
if avg_sla:
    exec_summary.append(f"Average SLA business time: {avg_sla}.")
executive_paragraph = ' '.join(exec_summary)

# Backup
if os.path.exists(docx_path):
    from shutil import copyfile
    copyfile(docx_path, backup_path)

# Open doc
doc = Document(docx_path)

# Helper replace function for cell or paragraph
def replace_text(obj_text, replacements):
    # obj_text is a string
    text = obj_text
    for placeholder, value in replacements.items():
        text = text.replace(placeholder, value)
    return text

# Build replacement mapping
replacements = {
    '{Total Cases}': str(total_cases),
    '{Open/UnResolved}': f"Open: {open_cases} / Resolved: {closed_cases}",
    '{TTR}': f"Median: {median_ttr if median_ttr else 'N/A'}; P90: {p90_ttr if p90_ttr else 'N/A'}",
    '{Top Trend}': trend_top,
    '{Executive summary}': executive_paragraph
}

# Chart mapping for phrases -> image path
chart_map = {
    '{CASES CREATED BY MONTH CHART}': monthly_chart,
    '{Top Products by Case Volume Chart}': product_chart,
    '{Priority Profile Chart}': priority_chart,
    '{Trend Distribution Chart}': trend_chart,
    '{Trend Movement Chart}': trend_movement_chart,
    '{Median vs P90 Chart}': median_vs_p90_chart,
    '{SLA Performance Chart}': sla_chart
}

# Replace in paragraphs
for para in doc.paragraphs:
    t = para.text
    newt = replace_text(t, replacements)
    if newt!=t:
        # clear runs then add
        for i in range(len(para.runs)-1,-1,-1):
            para.runs[i].clear()
        para.add_run(newt)

# Replace in tables and insert images where placeholders ask for charts
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            celltext = cell.text
            # check for chart placeholders
            for ph, img in chart_map.items():
                if ph in celltext and os.path.exists(img):
                    # clear cell
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.clear()
                    # insert image in first paragraph
                    cell.paragraphs[0].add_run().add_picture(img, width=Inches(3))
                    celltext = celltext.replace(ph, '')
            # generic replacements
            newcell = replace_text(celltext, replacements)
            if newcell!=celltext:
                # clear paragraphs
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.clear()
                cell.paragraphs[0].add_run(newcell)

# Save updated doc
final_path = os.path.join(out_dir,'Cases_Management_Report_Completed_filled.docx')
doc.save(final_path)
print('Saved filled report to', final_path)
print('Backup was saved to', backup_path)
