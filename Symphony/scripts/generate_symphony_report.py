import re
import sys
from collections import Counter
from pathlib import Path

import pandas as pd
from docx import Document

# Paths
BASE = Path(r"C:\Users\lee.booth\Documents\02_ServiceNow\Management_Reports")
CSV_PATH = BASE / "Symphony_Cases_With_SLA_Last_12-Months.csv"
TEMPLATE_PATH = BASE / "Symphony_Cases_Management_Report 2.docx"
OUTPUT_PATH = BASE / "Symphony_Cases_Management_Report_2_reworked.docx"

# Utilities
DUR_RE = re.compile(r"(?:(\d+)\s*day)|(?:(\d+)\s*days)|(?:(\d+)\s*hr)|(?:(\d+)\s*hrs)|(?:(\d+)\s*hour)|(?:(\d+)\s*hours)|(?:(\d+)\s*min)|(?:(\d+)\s*mins)|(?:(\d+)\s*minute)|(?:(\d+)\s*minutes)|(?:(\d+)\s*sec)|(?:(\d+)\s*secs)|(?:(\d+)\s*second)|(?:(\d+)\s*seconds)", re.IGNORECASE)

def parse_duration_to_seconds(s):
    if pd.isna(s) or s == "":
        return None
    # if numeric
    try:
        if isinstance(s, (int, float)):
            return int(s)
        s_str = str(s).strip()
        if s_str.isdigit():
            return int(s_str)
    except Exception:
        pass
    # parse human readable like "2 days, 11 hrs, 16 mins, 57 secs"
    days = hrs = mins = secs = 0
    m = re.search(r"(\d+)\s*day", s_str, re.IGNORECASE)
    if m:
        days = int(m.group(1))
    m = re.search(r"(\d+)\s*hr", s_str, re.IGNORECASE)
    if m:
        hrs = int(m.group(1))
    m = re.search(r"(\d+)\s*min", s_str, re.IGNORECASE)
    if m:
        mins = int(m.group(1))
    m = re.search(r"(\d+)\s*sec", s_str, re.IGNORECASE)
    if m:
        secs = int(m.group(1))
    total = days * 86400 + hrs * 3600 + mins * 60 + secs
    return total if total > 0 else 0


def seconds_to_human(seconds):
    if seconds is None or pd.isna(seconds):
        return ""
    seconds = int(seconds)
    parts = []
    days, rem = divmod(seconds, 86400)
    hrs, rem = divmod(rem, 3600)
    mins, secs = divmod(rem, 60)
    if days:
        parts.append(f"{days} day{'s' if days!=1 else ''}")
    if hrs:
        parts.append(f"{hrs} hr{'s' if hrs!=1 else ''}")
    if mins:
        parts.append(f"{mins} min{'s' if mins!=1 else ''}")
    if secs or not parts:
        parts.append(f"{secs} sec{'s' if secs!=1 else ''}")
    return ", ".join(parts)


# Read CSV
print(f"Reading {CSV_PATH}")
df = pd.read_csv(CSV_PATH, dtype=str)

# Normalize column names
df.columns = [c.strip() for c in df.columns]

# Keep Symphony product rows only
if 'product' in df.columns:
    df = df[df['product'].str.strip().str.lower() == 'symphony']

# Exclude PRBs: where u_problem contains PRB or the column seems to hold PRB refs
if 'u_problem' in df.columns:
    df = df[df['u_problem'].isna() | (df['u_problem'].str.strip() == '')]

# Parse dates
for col in ['sys_created_on', 'resolved_at']:
    if col in df.columns:
        df[col] = pd.to_datetime(df[col], dayfirst=True, errors='coerce')

# Parse SLA durations into seconds
for col in ['SLA_business_duration', 'SLA_duration']:
    if col in df.columns:
        df[col + '_secs'] = df[col].apply(parse_duration_to_seconds)

# Compute resolution seconds
if 'sys_created_on' in df.columns and 'resolved_at' in df.columns:
    df['resolution_secs'] = (df['resolved_at'] - df['sys_created_on']).dt.total_seconds()
else:
    df['resolution_secs'] = None

# Short descriptions: remove line feeds
if 'short_description' in df.columns:
    df['short_description_clean'] = df['short_description'].fillna('').str.replace('\n', ' ').str.replace('\r', ' ').str.strip()
else:
    df['short_description_clean'] = ''

# Totals
total_cases = len(df)
p4_count = df['priority'].fillna('').str.startswith('4').sum() if 'priority' in df.columns else 0

# Trends: top tokens
stopwords = set(["the","and","to","a","of","in","is","for","on","with","by","an","be","this","that","are","as","it","from","at","or","has","have"]) 
tokens = []
for s in df['short_description_clean']:
    words = re.findall(r"\w+", str(s).lower())
    words = [w for w in words if w not in stopwords and len(w) > 2]
    tokens.extend(words)
top_tokens = [t for t,_ in Counter(tokens).most_common(6)]

trends = []
for tok in top_tokens:
    mask = df['short_description_clean'].str.lower().str.contains(tok)
    sub = df[mask]
    if len(sub) == 0:
        continue
    trends.append({
        'token': tok,
        'count': len(sub),
        'avg_resolution_secs': int(sub['resolution_secs'].dropna().astype(float).mean()) if sub['resolution_secs'].dropna().size>0 else None,
        'avg_sla_business_secs': int(sub['SLA_business_duration_secs'].dropna().astype(float).mean()) if 'SLA_business_duration_secs' in sub.columns and sub['SLA_business_duration_secs'].dropna().size>0 else None,
        'avg_sla_actual_secs': int(sub['SLA_duration_secs'].dropna().astype(float).mean()) if 'SLA_duration_secs' in sub.columns and sub['SLA_duration_secs'].dropna().size>0 else None,
    })

# Overall averages
overall_avg_resolution = int(df['resolution_secs'].dropna().astype(float).mean()) if df['resolution_secs'].dropna().size>0 else None
overall_avg_sla_business = int(df['SLA_business_duration_secs'].dropna().astype(float).mean()) if 'SLA_business_duration_secs' in df.columns and df['SLA_business_duration_secs'].dropna().size>0 else None
overall_avg_sla_actual = int(df['SLA_duration_secs'].dropna().astype(float).mean()) if 'SLA_duration_secs' in df.columns and df['SLA_duration_secs'].dropna().size>0 else None

# Prepare recommendations (simple templates)
def recommend_for_token(tok):
    r = []
    r.append(f"Investigate root causes for '{tok}' trend and create a reusable knowledgebase article.")
    r.append("Provide targeted training or quick-reference guides to support teams.")
    r.append("Automate common fixes where possible and add proactive monitoring/alerts.")
    return r

# Build DOCX from template
doc = Document(TEMPLATE_PATH)
doc.add_page_break()
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(f'Total Symphony cases (last 12 months): {total_cases}')
doc.add_paragraph(f'P4 case count: {p4_count}')
if overall_avg_resolution:
    doc.add_paragraph(f'Average time to resolution: {seconds_to_human(overall_avg_resolution)}')
if overall_avg_sla_business:
    doc.add_paragraph(f'Average SLA business duration: {seconds_to_human(overall_avg_sla_business)}')
if overall_avg_sla_actual:
    doc.add_paragraph(f'Average SLA actual duration: {seconds_to_human(overall_avg_sla_actual)}')

doc.add_heading('Trends', level=1)
for t in trends:
    doc.add_heading(t['token'].capitalize(), level=2)
    doc.add_paragraph(f"Count: {t['count']}")
    doc.add_paragraph(f"Average resolution: {seconds_to_human(t['avg_resolution_secs'])}")
    doc.add_paragraph(f"Average SLA business: {seconds_to_human(t['avg_sla_business_secs'])}")
    doc.add_paragraph(f"Average SLA actual: {seconds_to_human(t['avg_sla_actual_secs'])}")
    doc.add_paragraph('Recommendations:')
    for r in recommend_for_token(t['token']):
        doc.add_paragraph(f"- {r}")

doc.add_heading('P4 Quick Wins', level=1)
p4 = df[df['priority'].fillna('').str.startswith('4')]
doc.add_paragraph(f'P4 count: {len(p4)}')
if len(p4)>0:
    # show top short descriptions in P4
    top_p4 = Counter(p4['short_description_clean'].fillna('')).most_common(5)
    doc.add_paragraph('Top P4 issue summaries:')
    for k,cnt in top_p4:
        doc.add_paragraph(f'- {k} ({cnt} cases)')
    doc.add_paragraph('Quick-win recommendations:')
    doc.add_paragraph('- Triage P4s into actionable fixes and document straightforward remediation steps.')
    doc.add_paragraph('- Implement monitoring/alerts to catch regressions earlier.')

doc.save(OUTPUT_PATH)
print(f"Report written to {OUTPUT_PATH}")
