from docx import Document
from pathlib import Path

TEMPLATE = Path(r"C:\Users\lee.booth\Documents\02_ServiceNow\Management_Reports\Symphony\templates\Symphony_Cases_Management_Report_Template.docx")
if not TEMPLATE.exists():
    print('Template not found:', TEMPLATE)
    raise SystemExit(1)

doc = Document(TEMPLATE)

print('Paragraph headings:')
for i,p in enumerate(doc.paragraphs[:200]):
    style = p.style.name if p.style else ''
    if style.lower().startswith('heading') and p.text.strip():
        print(f'  {i}: {style} -> "{p.text.strip()}"')

print('\nTables:')
for ti, table in enumerate(doc.tables):
    print(f' Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols')
    # print first 3 rows
    maxr = min(3, len(table.rows))
    for r in range(maxr):
        cells = [table.rows[r].cells[c].text.strip().replace('\n',' ') for c in range(len(table.columns))]
        print('   Row', r, ' | '.join(cells))
    print('  ...')
