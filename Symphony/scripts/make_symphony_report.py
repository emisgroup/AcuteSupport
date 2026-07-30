import os
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

folder = r"C:\Users\lee.booth\Documents\02_ServiceNow\Management_Reports"
bak = os.path.join(folder, 'Symphony_Cases_With_SLA_Last_12-Months.csv.bak')
if not os.path.exists(bak):
    raise SystemExit('Backup CSV not found: '+bak)

# Read combined bak (numeric SLA columns)
df = pd.read_csv(bak, dtype=str, low_memory=False)
# ensure datetime
if 'sys_created_on' in df.columns:
    df['sys_created_on_dt'] = pd.to_datetime(df['sys_created_on'], errors='coerce')
    df['month'] = df['sys_created_on_dt'].dt.to_period('M').astype(str)
else:
    df['month'] = 'unknown'

# SLA numeric columns might be named 'SLA_business_duration' and 'SLA_duration'
for col in ['SLA_business_duration','SLA_duration']:
    if col in df.columns:
        df[col+'_secs'] = pd.to_numeric(df[col], errors='coerce')
    else:
        df[col+'_secs'] = pd.NA

# Monthly trend
monthly = df.groupby('month').size().reset_index(name='count').sort_values('month')
plt.figure(figsize=(8,3))
plt.plot(monthly['month'], monthly['count'], marker='o')
plt.xticks(rotation=45)
plt.title('Monthly Case Trend')
plt.tight_layout()
trend_img = os.path.join(folder, 'sym_monthly_trend.png')
plt.savefig(trend_img)
plt.close()

# Averages
avg_business = df['SLA_business_duration_secs'].dropna().astype(float).mean()
avg_sla = df['SLA_duration_secs'].dropna().astype(float).mean()

# Top problems
top_probs = pd.read_csv(os.path.join(folder, 'Symphony_top_problems.csv')) if os.path.exists(os.path.join(folder, 'Symphony_top_problems.csv')) else None

# P4 cases
if 'priority' in df.columns:
    p4_mask = df['priority'].str.contains('\\b4\\b', na=False, case=False) | df['priority'].str.upper().str.contains('P4', na=False)
else:
    p4_mask = pd.Series([False]*len(df))
df_p4 = df[p4_mask]
p4_count = len(df_p4)
p4_avg_business = df_p4['SLA_business_duration_secs'].dropna().astype(float).mean()
p4_avg_sla = df_p4['SLA_duration_secs'].dropna().astype(float).mean()

# Quick wins: P4 problems with high count but low avg business duration
quick_wins = []
if top_probs is not None:
    merged = top_probs.merge(df[['u_problem']], left_on='Problem', right_on='u_problem', how='left')
# Instead compute quick-wins by problem in df_p4
p4_by_problem = df_p4.groupby('u_problem').agg(count=('number','count'), avg_bus=('SLA_business_duration_secs', 'mean'), avg_sla=('SLA_duration_secs','mean')).reset_index().sort_values('count', ascending=False)
for _, row in p4_by_problem.head(5).iterrows():
    if pd.notna(row['avg_bus']) and row['avg_bus'] < 86400*3:  # less than 3 days
        quick_wins.append({'problem': row['u_problem'], 'count': int(row['count']), 'avg_business_days': float(row['avg_bus'])/86400})

# Build DOCX
doc = Document()
doc.add_heading('Symphony Management Report', level=1)
# Executive summary
doc.add_heading('Executive Summary', level=2)
doc.add_paragraph(f'Total cases (last 12 months): {len(df):,}')
doc.add_paragraph(f'Average business duration: {avg_business:.2f} seconds ({(avg_business or 0)/86400:.2f} days)')
doc.add_paragraph(f'Average SLA duration: {avg_sla:.2f} seconds ({(avg_sla or 0)/86400:.2f} days)')

# Trend chart
doc.add_heading('Monthly Trends', level=2)
doc.add_paragraph('Monthly case counts are shown below:')
doc.add_picture(trend_img, width=Inches(6))

# Averages and SLA timings
doc.add_heading('Average Timings', level=2)

t = doc.add_paragraph()
t.add_run('Average business duration: ').bold = True
t.add_run(f'{avg_business:.0f} secs ({(avg_business or 0)/86400:.2f} days)')

t2 = doc.add_paragraph()
t2.add_run('Average SLA duration: ').bold = True
t2.add_run(f'{avg_sla:.0f} secs ({(avg_sla or 0)/86400:.2f} days)')

# Top problems
doc.add_heading('Top Problems', level=2)
if top_probs is not None:
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text='Problem'; hdr[1].text='Count'; hdr[2].text='AvgBusinessSecs'; hdr[3].text='AvgSLAsecs'
    for _, r in top_probs.iterrows():
        row = table.add_row().cells
        row[0].text = str(r['Problem'])
        row[1].text = str(r['Count'])
        row[2].text = str(r.get('AvgBusiness', ''))
        row[3].text = str(r.get('AvgSLA', ''))
else:
    doc.add_paragraph('No top problems CSV available.')

# P4 section
doc.add_heading('P4 Cases', level=2)
doc.add_paragraph(f'P4 case count: {p4_count}')
doc.add_paragraph(f'P4 average business duration: {p4_avg_business or 0:.0f} secs ({(p4_avg_business or 0)/86400:.2f} days)')
doc.add_paragraph(f'P4 average SLA duration: {p4_avg_sla or 0:.0f} secs ({(p4_avg_sla or 0)/86400:.2f} days)')

# Quick wins
doc.add_heading('Quick Wins (P4)', level=3)
if quick_wins:
    for qw in quick_wins:
        doc.add_paragraph(f"Problem: {qw['problem']} — Count: {qw['count']} — Avg business days: {qw['avg_business_days']:.2f}")
    doc.add_paragraph('Recommended quick wins: create standard runbooks, knowledge base articles, automation for common fixes, and targeted training for support teams to reduce repeat P4 cases.')
else:
    doc.add_paragraph('No clear quick-wins identified in P4 top problems.')

# Management recommendations
doc.add_heading('Management Recommendations', level=2)
doc.add_paragraph('1. Knowledge base and self-service: Create KB articles for top recurring problems to deflect simple P4 issues.')
doc.add_paragraph('2. Automation: Implement automated remediation for repetitive tasks (restarts, cache clears).')
doc.add_paragraph('3. Triage and routing: Improve initial triage to assign problems to appropriate resolver groups earlier.')
doc.add_paragraph('4. Monitoring and alerts: Add proactive monitoring for top problem indicators to reduce time-to-detect.')

doc.add_paragraph('5. Process: Regularly review top problem trends and measure effectiveness of KB/automation with monthly KPIs.')

# Save
outdoc = os.path.join(folder, 'Symphony_Management_Report.docx')
doc.save(outdoc)
# Cleanup
try:
    os.remove(trend_img)
except:
    pass
print('WROTE', outdoc)
