import re
from pathlib import Path
from collections import Counter
import pandas as pd
from docx import Document
from datetime import datetime
import matplotlib.pyplot as plt
from docx.shared import Inches

BASE = Path(r"C:\Users\lee.booth\Documents\02_ServiceNow\Management_Reports")
SYM_DIR = BASE / 'Symphony'
RAW = SYM_DIR / 'raw'
SCRIPTS = SYM_DIR / 'scripts'
REPORTS = SYM_DIR / 'reports'
TEMPLATES = SYM_DIR / 'templates'

CSV_FILE = RAW / 'Symphony_Cases_With_SLA_Last_12-Months.csv'
TEMPLATE = TEMPLATES / 'Symphony_Cases_Management_Report_Template.docx'

OUT_NAME = f"Symphony_Report_Filled_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
OUT_PATH = REPORTS / OUT_NAME

# Helpers for durations
re_day = re.compile(r"(\d+)\s*day", re.IGNORECASE)
re_hr = re.compile(r"(\d+)\s*hr", re.IGNORECASE)
re_min = re.compile(r"(\d+)\s*min", re.IGNORECASE)
re_sec = re.compile(r"(\d+)\s*sec", re.IGNORECASE)

STOPWORDS = set(["the","and","to","a","of","in","is","for","on","with","by","an","be","this","that","are","as","it","from","at","or","has","have"])


def parse_duration_to_seconds(s):
    if pd.isna(s) or s == "":
        return None
    try:
        if isinstance(s, (int,float)):
            return int(s)
        s_str = str(s)
        if s_str.isdigit():
            return int(s_str)
    except Exception:
        pass
    days = hrs = mins = secs = 0
    m = re_day.search(s_str)
    if m: days = int(m.group(1))
    m = re_hr.search(s_str)
    if m: hrs = int(m.group(1))
    m = re_min.search(s_str)
    if m: mins = int(m.group(1))
    m = re_sec.search(s_str)
    if m: secs = int(m.group(1))
    total = days*86400 + hrs*3600 + mins*60 + secs
    return total if total>0 else 0


def seconds_to_human(seconds):
    if seconds is None or pd.isna(seconds):
        return ""
    seconds = int(seconds)
    parts = []
    days, rem = divmod(seconds, 86400)
    hrs, rem = divmod(rem, 3600)
    mins, secs = divmod(rem, 60)
    if days: parts.append(f"{days} day{'s' if days!=1 else ''}")
    if hrs: parts.append(f"{hrs} hr{'s' if hrs!=1 else ''}")
    if mins: parts.append(f"{mins} min{'s' if mins!=1 else ''}")
    if secs or not parts: parts.append(f"{secs} sec{'s' if secs!=1 else ''}")
    return ", ".join(parts)

# Read template
if not TEMPLATE.exists():
    raise SystemExit(f"Template not found: {TEMPLATE}")

doc = Document(TEMPLATE)

# Read CSV
if not CSV_FILE.exists():
    raise SystemExit(f"CSV not found: {CSV_FILE}")

df = pd.read_csv(CSV_FILE, dtype=str)
# normalize
df.columns = [c.strip() for c in df.columns]
# keep Symphony only
if 'product' in df.columns:
    df = df[df['product'].fillna('').str.strip().str.lower()=='symphony']
# remove PRBs
if 'u_problem' in df.columns:
    df = df[df['u_problem'].isna() | (df['u_problem'].str.strip()=='')]
# clean short desc
if 'short_description' in df.columns:
    df['short_description_clean'] = df['short_description'].fillna('').str.replace('\n',' ').str.replace('\r',' ').str.strip()
else:
    df['short_description_clean'] = ''
# parse dates
for col in ['sys_created_on','resolved_at']:
    if col in df.columns:
        df[col] = pd.to_datetime(df[col], dayfirst=True, errors='coerce')
# resolution
if 'sys_created_on' in df.columns and 'resolved_at' in df.columns:
    df['resolution_secs'] = (df['resolved_at'] - df['sys_created_on']).dt.total_seconds()
else:
    df['resolution_secs'] = None
# parse SLA durations
for col in ['SLA_business_duration','SLA_duration']:
    if col in df.columns:
        df[col + '_secs'] = df[col].apply(parse_duration_to_seconds)

# Metrics
total_cases = len(df)
open_count = int(df[~df['state'].fillna('').str.lower().isin(['closed','resolved'])].shape[0]) if 'state' in df.columns else 0
p4_count = int(df['priority'].fillna('').str.startswith('4').sum()) if 'priority' in df.columns else 0
overall_avg_resolution = int(df['resolution_secs'].dropna().astype(float).mean()) if df['resolution_secs'].dropna().size>0 else None
median_resolution = int(df['resolution_secs'].dropna().astype(float).median()) if df['resolution_secs'].dropna().size>0 else None
overall_avg_sla_business = int(df['SLA_business_duration_secs'].dropna().astype(float).mean()) if 'SLA_business_duration_secs' in df.columns and df['SLA_business_duration_secs'].dropna().size>0 else None
overall_avg_sla_actual = int(df['SLA_duration_secs'].dropna().astype(float).mean()) if 'SLA_duration_secs' in df.columns and df['SLA_duration_secs'].dropna().size>0 else None

# Top themes from short description
tokens = []
for s in df['short_description_clean']:
    words = re.findall(r"\w+", str(s).lower())
    words = [w for w in words if w not in STOPWORDS and len(w)>2]
    tokens.extend(words)
common = Counter(tokens).most_common(10)

# Top accounts
top_accounts = Counter(df['account'].fillna('Unknown')).most_common(10)

# Priority distribution
priority_counts = df['priority'].fillna('Unknown').value_counts()
# State counts
state_counts = df['state'].fillna('Unknown').value_counts()

# Monthly trend
if 'sys_created_on' in df.columns:
    df['month'] = df['sys_created_on'].dt.to_period('M')
    monthly = df.groupby('month').size().sort_index()
else:
    monthly = pd.Series([], dtype=int)

# Start filling template tables by index (observed structure)
# Table 0: Executive summary key cells row 1 (index 1)
try:
    t0 = doc.tables[0]
    # Row 1 has labels: Total cases | Open / unresolved | Median time to resolution | Largest product stream
    if t0.rows and len(t0.rows) > 1:
        r = t0.rows[1]
        r.cells[0].text = str(total_cases)
        r.cells[1].text = str(open_count)
        r.cells[2].text = seconds_to_human(median_resolution)
        # largest product stream -> use top theme
        r.cells[3].text = common[0][0] if common else 'Symphony'
except IndexError:
    pass

# Table 2: metrics table (Metric|Value|Notes) -- fill rows 1.. up to available
metrics = [
    ("Total cases", str(total_cases), "Last 12 months"),
    ("Open / unresolved", str(open_count), "Cases not closed or resolved"),
    ("P4 cases", str(p4_count), "Priority 4"),
    ("Average time to resolution", seconds_to_human(overall_avg_resolution), "Mean of resolved cases"),
    ("Median time to resolution", seconds_to_human(median_resolution), "Median of resolved cases"),
    ("Average SLA business timing", seconds_to_human(overall_avg_sla_business), "SLA business duration mean"),
    ("Average SLA actual (SLA_Duration)", seconds_to_human(overall_avg_sla_actual), "SLA elapsed time mean"),
]
try:
    t2 = doc.tables[2]
    max_rows = len(t2.rows)
    for i, m in enumerate(metrics, start=1):
        if i >= max_rows:
            break
        t2.rows[i].cells[0].text = m[0]
        t2.rows[i].cells[1].text = m[1]
        t2.rows[i].cells[2].text = m[2]
except IndexError:
    pass

# Table 3: chart placeholders (2 cells) -> insert monthly chart and priority pie
chart1 = REPORTS / 'chart_monthly.png'
chart2 = REPORTS / 'chart_priority.png'
# generate charts
REPORTS.mkdir(parents=True, exist_ok=True)
if not monthly.empty:
    plt.figure(figsize=(6,3))
    monthly.index = monthly.index.to_timestamp()
    plt.plot(monthly.index, monthly.values, marker='o')
    plt.title('Monthly case volume')
    plt.ylabel('Cases')
    plt.tight_layout()
    plt.savefig(chart1)
    plt.close()
# priority chart
plt.figure(figsize=(4,3))
priority_counts.plot(kind='bar')
plt.title('Priority distribution')
plt.tight_layout()
plt.savefig(chart2)
plt.close()

try:
    t3 = doc.tables[3]
    # left cell -> monthly
    if chart1.exists():
        t3.rows[0].cells[0].paragraphs[0].clear()
        run = t3.rows[0].cells[0].paragraphs[0].add_run()
        run.add_picture(str(chart1), width=Inches(3))
    if chart2.exists():
        t3.rows[0].cells[1].paragraphs[0].clear()
        run = t3.rows[0].cells[1].paragraphs[0].add_run()
        run.add_picture(str(chart2), width=Inches(3))
except IndexError:
    pass

# Table 4: Rank | Product | Cases -> use top themes as 'product streams'
try:
    t4 = doc.tables[4]
    for i, (tok, cnt) in enumerate(common[:len(t4.rows)-1], start=1):
        if i >= len(t4.rows): break
        t4.rows[i].cells[0].text = str(i)
        t4.rows[i].cells[1].text = tok.capitalize()
        t4.rows[i].cells[2].text = str(cnt)
except IndexError:
    pass

# Table 6: State | Cases | Priority | Cases
try:
    t6 = doc.tables[6]
    # fill states in first two columns
    state_items = list(state_counts.items())
    prio_items = list(priority_counts.items())
    maxr = len(t6.rows)-1
    for i in range(maxr):
        row = t6.rows[i+1]
        if i < len(state_items):
            row.cells[0].text = state_items[i][0]
            row.cells[1].text = str(state_items[i][1])
        if i < len(prio_items):
            row.cells[2].text = prio_items[i][0]
            row.cells[3].text = str(prio_items[i][1])
except IndexError:
    pass

# Table 8: Themes (Theme | Matching cases | Recent signal | Management implication)
try:
    t8 = doc.tables[8]
    for i, (tok, cnt) in enumerate(common[:len(t8.rows)-1], start=1):
        if i >= len(t8.rows): break
        row = t8.rows[i]
        row.cells[0].text = tok.capitalize()
        row.cells[1].text = str(cnt)
        # recent signal: cases since June of current year
        year = datetime.now().year
        june1 = pd.Timestamp(year=year, month=6, day=1)
        recent = df[(df['sys_created_on']>=june1) & (df['short_description_clean'].str.lower().str.contains(tok))] if 'sys_created_on' in df.columns else df[df['short_description_clean'].str.lower().str.contains(tok)]
        row.cells[2].text = str(len(recent))
        row.cells[3].text = 'Create KB article and run targeted training'[:200]
except IndexError:
    pass

# Table 9: Account concentration
try:
    t9 = doc.tables[9]
    for i, (acc, cnt) in enumerate(top_accounts[:len(t9.rows)-1], start=1):
        if i >= len(t9.rows): break
        row = t9.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = acc
        row.cells[2].text = str(cnt)
        # open count for that account
        open_ct = int(df[(df['account']==acc) & (~df['state'].fillna('').str.lower().isin(['closed','resolved']))].shape[0]) if 'state' in df.columns else 0
        row.cells[3].text = str(open_ct)
except IndexError:
    pass

# Table 10: Recommendations
try:
    t10 = doc.tables[10]
    recs = [
        ("Knowledgebase articles","Document fixes for top themes and link from ticket triage."),
        ("Training","Deliver quick sessions for support on top 3 themes."),
        ("Automation","Script routine fixes and add runbooks."),
        ("Monitoring","Add alerts for repeat failures and key metrics."),
        ("Triage","Introduce a P4 quick-win path to resolve in <4 hrs where possible."),
    ]
    for i, (rname, rationale) in enumerate(recs[:len(t10.rows)-1], start=1):
        row = t10.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = rname
        row.cells[2].text = rationale
except IndexError:
    pass

# Save doc
REPORTS.mkdir(parents=True, exist_ok=True)
doc.save(OUT_PATH)
print(f"Filled report written to: {OUT_PATH}")
print('Charts written to:', chart1, chart2)
