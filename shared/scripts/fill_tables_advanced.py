from docx import Document
from docx.shared import Inches
import os
import pandas as pd
from shutil import copyfile

# Base directories (dynamic)
BASE_DIR = os.getcwd()
RAW_DIR = os.path.join(BASE_DIR, 'data', 'raw')
PROCESSED_DIR = os.path.join(BASE_DIR, 'data', 'processed')
tables_dir = os.path.join(BASE_DIR, 'outputs', 'tables')
reports_dir = os.path.join(BASE_DIR, 'outputs', 'reports')

filled_doc = os.path.join(reports_dir, 'Cases_Management_Report_Completed_filled.docx')
if not os.path.exists(filled_doc):
    filled_doc = os.path.join(reports_dir, 'Cases_Management_Report_Completed.docx')

backup_doc = os.path.join(reports_dir, 'Cases_Management_Report_Completed_filled_backup.docx')

merged_csv = os.path.join(PROCESSED_DIR, 'Merged_Cases_With_SLA_Formatted.csv')
if not os.path.exists(merged_csv):
    merged_csv = os.path.join(RAW_DIR, 'Merged_Cases_With_SLA_Formatted.csv')

trend_dist_csv = os.path.join(tables_dir, 'trend_distribution.csv')
product_csv = os.path.join(tables_dir, 'product_distribution.csv')
monthly_csv = os.path.join(tables_dir, 'monthly_case_trend.csv')
priority_csv = os.path.join(tables_dir, 'priority_profile.csv')
trend_movement_csv = os.path.join(tables_dir, 'trend_movement.csv')
account_csv = os.path.join(tables_dir, 'account_distribution.csv')

# load data
if os.path.exists(merged_csv):
    df = pd.read_csv(merged_csv, encoding='utf-8')
else:
    raise SystemExit('Merged CSV not found')

trend_dist = pd.read_csv(trend_dist_csv) if os.path.exists(trend_dist_csv) else pd.DataFrame(columns=['trend','count'])
product_dist = pd.read_csv(product_csv) if os.path.exists(product_csv) else pd.DataFrame(columns=['product','count'])
monthly = pd.read_csv(monthly_csv) if os.path.exists(monthly_csv) else pd.DataFrame(columns=['created_month','cases'])
priority = pd.read_csv(priority_csv, index_col=0)['count'] if os.path.exists(priority_csv) else pd.Series(dtype=int)
trend_movement = pd.read_csv(trend_movement_csv, index_col=0) if os.path.exists(trend_movement_csv) else pd.DataFrame()

# compute account distribution
if 'account' in df.columns:
    acc = df.groupby('account').size().reset_index(name='cases').sort_values('cases',ascending=False)
    # open count
    open_mask = df['state'].str.contains('Closed', na=False)==False if 'state' in df.columns else pd.Series([False]*len(df))
    open_counts = df[open_mask].groupby('account').size().reset_index(name='open') if len(df)>0 else pd.DataFrame()
    account_df = acc.merge(open_counts, on='account', how='left').fillna(0)
    account_df['open'] = account_df['open'].astype(int)
else:
    account_df = pd.DataFrame(columns=['account','cases','open'])

# prepare trend recent signal from June onwards (2026-06)
# trend_movement has created_month as index (string), columns are trends
recent_threshold = '2026-06'
recent_counts = {}
if not trend_movement.empty:
    # ensure index is string
    tm = trend_movement.copy()
    tm.index = tm.index.astype(str)
    for col in tm.columns:
        try:
            s = tm.loc[[c for c in tm.index if c>=recent_threshold], col].sum()
        except Exception:
            s = 0
        recent_counts[col] = int(s)

# management implication: simple rule-based
implications = {}
total = len(df)
for _,row in trend_dist.iterrows():
    t = row['trend']
    cnt = int(row['count'])
    pct = cnt/total if total>0 else 0
    if pct > 0.1:
        imp = 'High volume — consider KB and automation.'
    elif pct > 0.03:
        imp = 'Medium volume — review for triage improvements.'
    else:
        imp = 'Low volume — monitor; consider knowledge article if recurring.'
    # if recent spike
    rc = recent_counts.get(t,0)
    if rc>0 and rc/ max(1,cnt) > 0.3:
        imp += ' Recent uptick observed.'
    implications[t]=imp

# open doc and backup
if os.path.exists(filled_doc):
    copyfile(filled_doc, backup_doc)
    doc = Document(filled_doc)
else:
    raise SystemExit('Filled DOCX not found')

# Helper to clear cell
def clear_cell(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.clear()

# Table indices based on inspection
# Table 2 (index 2): Metric table -> fill Total cases analysed
if len(doc.tables) > 2:
    t2 = doc.tables[2]
    # find row where first cell contains 'Total cases analysed' (case-insensitive) else use row1
    row_idx = None
    for i,row in enumerate(t2.rows):
        if 'total cases analysed' in row.cells[0].text.lower():
            row_idx = i
            break
    if row_idx is None:
        row_idx = 1
    clear_cell(t2.rows[row_idx].cells[1])
    t2.rows[row_idx].cells[1].paragraphs[0].add_run(str(total))

# Table 4 (index 4): Product ranking
if len(doc.tables) > 4:
    t4 = doc.tables[4]
    # clear existing content rows, then insert top products
    # keep header row, remove other rows
    while len(t4.rows) > 1:
        t4._tbl.remove(t4.rows[-1]._tr)
    for i, (prod, cnt) in enumerate(product_dist.head(10).values.tolist(), start=1):
        row = t4.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = str(prod)
        row.cells[2].text = str(int(cnt))

# Table 6 (index 6): State and Priority counts
if len(doc.tables) > 6:
    t6 = doc.tables[6]
    # compute state counts
    state_counts = df['state'].fillna('Unknown').value_counts()
    pr_counts = priority if not priority.empty else df['priority'].fillna('Unknown').value_counts()
    maxrows = max(len(state_counts), len(pr_counts))
    # remove existing data rows (keep header)
    while len(t6.rows) > 1:
        t6._tbl.remove(t6.rows[-1]._tr)
    for i in range(maxrows):
        sr = t6.add_row()
        sname = state_counts.index[i] if i < len(state_counts) else ''
        scnt = state_counts.iloc[i] if i < len(state_counts) else ''
        pname = pr_counts.index[i] if i < len(pr_counts) else ''
        pcnt = int(pr_counts.iloc[i]) if i < len(pr_counts) else ''
        # Write safely according to number of columns in the table row (template may vary)
        vals = [sname, scnt, pname, pcnt]
        ncols = len(sr.cells)
        for j in range(min(ncols, 4)):
            # clear existing runs
            for p in sr.cells[j].paragraphs:
                for r in p.runs:
                    r.clear()
            sr.cells[j].paragraphs[0].add_run(str(vals[j]))

# Table 8 (index 8): Theme table - populate trends and metrics
if len(doc.tables) > 8:
    t8 = doc.tables[8]
    # remove existing data rows except header
    while len(t8.rows) > 1:
        t8._tbl.remove(t8.rows[-1]._tr)
    # For each trend in trend_dist, add row: Theme | Matching cases | Recent signal | Management implication
    for _, row in trend_dist.iterrows():
        name = row['trend']
        cnt = int(row['count'])
        recent = recent_counts.get(name, 0)
        impl = implications.get(name, '')
        newr = t8.add_row()
        newr.cells[0].text = str(name)
        newr.cells[1].text = str(cnt)
        newr.cells[2].text = str(recent)
        newr.cells[3].text = impl

# Table 9 (index 9): Account ranking
if len(doc.tables) > 9:
    t9 = doc.tables[9]
    while len(t9.rows) > 1:
        t9._tbl.remove(t9.rows[-1]._tr)
    for i, r in enumerate(account_df.head(10).itertuples(), start=1):
        nr = t9.add_row()
        vals = [str(i), str(r.account), str(int(r.cases)), str(int(r.open))]
        ncols = len(nr.cells)
        for j in range(min(ncols, 4)):
            for p in nr.cells[j].paragraphs:
                for run in p.runs:
                    run.clear()
            nr.cells[j].paragraphs[0].add_run(vals[j])

# Save updated doc
out_path = os.path.join(reports_dir, 'Cases_Management_Report_Completed_tables_filled.docx')
doc.save(out_path)
print('Saved filled tables DOCX to', out_path)
print('Backup saved to', backup_doc)
